import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

target_dir = "/Users/howardliao/Desktop/Howard/Howard_CISO"
photo_path = "/Users/howardliao/Library/Application Support/Hermes/composer-images/composer_2026-08-21_07-55-16-061_53c1dd.jpg"
asset_dir = "/Users/howardliao/Desktop/Howard/Howard_CISO/assets"

COLOR_PRIMARY = RGBColor(15, 41, 66)      # #0F2942 Deep Navy
COLOR_SECONDARY = RGBColor(30, 58, 138)  # #1E3A8A Executive Blue
COLOR_SLATE = RGBColor(51, 65, 85)        # #334155 Slate
COLOR_TEXT = RGBColor(31, 41, 55)         # #1F2937 Charcoal Body Text
COLOR_MUTED = RGBColor(100, 116, 139)     # #64748B Subtitle/Meta Gray
COLOR_LINK = RGBColor(37, 99, 235)        # #2563EB Blue Link
FONT_FAMILY = 'Arial'

def create_base_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    return doc

def set_run_font(run, name=FONT_FAMILY, size=Pt(10), bold=False, italic=False, color=COLOR_TEXT, east_asia="Microsoft JhengHei"):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{east_asia}" w:cs="{name}"/>')
    rPr.append(rFonts)

def add_heading_1(doc, text, space_before=14, space_after=4, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=Pt(13), bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="12" w:space="3" w:color="1E3A8A"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_subheading(doc, text, space_before=8, space_after=2, color=COLOR_SECONDARY, size=Pt(10.5), east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color, east_asia=east_asia)
    return p

def add_body_p(doc, text, space_before=0, space_after=4, line_spacing=1.15, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    set_run_font(run, size=Pt(9.5), color=COLOR_TEXT, east_asia=east_asia)
    return p

def add_bullet(doc, text, space_before=0, space_after=2, line_spacing=1.15, bold_prefix_colon=True, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.left_indent = Inches(0.22)
    
    if bold_prefix_colon and (':' in text or '：' in text) and not text.startswith('http'):
        sep = ':' if ':' in text else '：'
        parts = text.split(sep, 1)
        r1 = p.add_run(parts[0] + sep)
        set_run_font(r1, size=Pt(9.5), bold=True, color=COLOR_TEXT, east_asia=east_asia)
        r2 = p.add_run(parts[1])
        set_run_font(r2, size=Pt(9.5), color=COLOR_TEXT, east_asia=east_asia)
    else:
        run = p.add_run(text)
        set_run_font(run, size=Pt(9.5), color=COLOR_TEXT, east_asia=east_asia)
    return p

def add_header(doc, name_text, title_text, subtitle_text, contact_text1, contact_text2, east_asia="Microsoft JhengHei"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_left = table.columns[0]
    col_right = table.columns[1]
    col_left.width = Inches(1.5)
    col_right.width = Inches(5.5)

    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    cell_left.width = Inches(1.5)
    cell_right.width = Inches(5.5)

    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for cell in [cell_left, cell_right]:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="none"/>'
            '<w:left w:val="none"/>'
            '<w:bottom w:val="none"/>'
            '<w:right w:val="none"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    p_photo = cell_left.paragraphs[0]
    p_photo.paragraph_format.space_before = Pt(0)
    p_photo.paragraph_format.space_after = Pt(0)
    if os.path.exists(photo_path):
        run_photo = p_photo.add_run()
        run_photo.add_picture(photo_path, width=Inches(1.35))

    p_name = cell_right.paragraphs[0]
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run(name_text)
    set_run_font(run_name, size=Pt(19), bold=True, color=COLOR_PRIMARY, east_asia=east_asia)

    p_title = cell_right.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(1)
    run_t1 = p_title.add_run(title_text)
    set_run_font(run_t1, size=Pt(11), bold=True, color=COLOR_SECONDARY, east_asia=east_asia)

    p_sub = cell_right.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(4)
    run_t2 = p_sub.add_run(subtitle_text)
    set_run_font(run_t2, size=Pt(10), bold=True, color=COLOR_SLATE, east_asia=east_asia)

    p_c1 = cell_right.add_paragraph()
    p_c1.paragraph_format.space_before = Pt(0)
    p_c1.paragraph_format.space_after = Pt(1)
    r_loc = p_c1.add_run(contact_text1)
    set_run_font(r_loc, size=Pt(9), color=COLOR_SLATE, east_asia=east_asia)

    p_c2 = cell_right.add_paragraph()
    p_c2.paragraph_format.space_before = Pt(0)
    p_c2.paragraph_format.space_after = Pt(0)
    r_c2 = p_c2.add_run(contact_text2)
    set_run_font(r_c2, size=Pt(9), color=COLOR_SLATE, east_asia=east_asia)

    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after = Pt(0)

def add_exp_item(doc, role, group_note, company_line, period, desc, scopes, achs, leaving=None, add_ctx=None, east_asia="Microsoft JhengHei", labels=None):
    if labels is None:
        labels = {"scope": "Leadership Scope", "ach": "Selected Achievements & Impact", "ctx": "Additional Context", "leaving": "Reason for Leaving"}
    
    p_r = doc.add_paragraph()
    p_r.paragraph_format.space_before = Pt(10)
    p_r.paragraph_format.space_after = Pt(1)
    p_r.paragraph_format.keep_with_next = True
    r1 = p_r.add_run(role)
    set_run_font(r1, size=Pt(11.5), bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    if group_note:
        p_cn = doc.add_paragraph()
        p_cn.paragraph_format.space_before = Pt(0)
        p_cn.paragraph_format.space_after = Pt(1)
        p_cn.paragraph_format.keep_with_next = True
        rcn = p_cn.add_run(group_note)
        set_run_font(rcn, size=Pt(9.5), color=COLOR_SLATE, east_asia=east_asia)
        
    p_m = doc.add_paragraph()
    p_m.paragraph_format.space_before = Pt(0)
    p_m.paragraph_format.space_after = Pt(2)
    p_m.paragraph_format.keep_with_next = True
    rm = p_m.add_run(f"{company_line}\n{period}")
    set_run_font(rm, size=Pt(9.5), bold=True, color=COLOR_SECONDARY, east_asia=east_asia)
    
    if desc:
        p_sd = doc.add_paragraph()
        p_sd.paragraph_format.space_before = Pt(2)
        p_sd.paragraph_format.space_after = Pt(3)
        p_sd.paragraph_format.keep_with_next = True
        rsd = p_sd.add_run(desc)
        set_run_font(rsd, size=Pt(9.5), italic=True, color=COLOR_TEXT, east_asia=east_asia)
        
    if scopes:
        p_sh = doc.add_paragraph()
        p_sh.paragraph_format.space_before = Pt(3)
        p_sh.paragraph_format.space_after = Pt(1)
        p_sh.paragraph_format.keep_with_next = True
        rsh = p_sh.add_run(labels["scope"])
        set_run_font(rsh, size=Pt(9.5), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        for s in scopes:
            add_bullet(doc, s, space_after=1.5, bold_prefix_colon=False, east_asia=east_asia)
            
    if achs:
        p_ah = doc.add_paragraph()
        p_ah.paragraph_format.space_before = Pt(3)
        p_ah.paragraph_format.space_after = Pt(1)
        p_ah.paragraph_format.keep_with_next = True
        rah = p_ah.add_run(labels["ach"])
        set_run_font(rah, size=Pt(9.5), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        for a in achs:
            add_bullet(doc, a, space_after=2, bold_prefix_colon=False, east_asia=east_asia)
            
    if add_ctx:
        p_ch = doc.add_paragraph()
        p_ch.paragraph_format.space_before = Pt(3)
        p_ch.paragraph_format.space_after = Pt(1)
        p_ch.paragraph_format.keep_with_next = True
        rch = p_ch.add_run(labels["ctx"])
        set_run_font(rch, size=Pt(9.5), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        add_body_p(doc, add_ctx, space_after=3, east_asia=east_asia)
        
    if leaving:
        p_lh = doc.add_paragraph()
        p_lh.paragraph_format.space_before = Pt(3)
        p_lh.paragraph_format.space_after = Pt(1)
        p_lh.paragraph_format.keep_with_next = True
        rlh = p_lh.add_run(labels["leaving"])
        set_run_font(rlh, size=Pt(9.5), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after = Pt(4)
        rl = p_l.add_run(leaving)
        set_run_font(rl, size=Pt(9), italic=True, color=COLOR_MUTED, east_asia=east_asia)

def add_media_entry(doc, title, org_date, desc, url=None, image_filename=None, img_width=Inches(3.2), east_asia="Microsoft JhengHei", url_label="🔗 Verified Source / DOI / URL: "):
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(7)
    p_t.paragraph_format.space_after = Pt(1)
    p_t.paragraph_format.keep_with_next = True
    r_t = p_t.add_run(title)
    set_run_font(r_t, size=Pt(10.5), bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(2)
    p_meta.paragraph_format.keep_with_next = True
    r_m = p_meta.add_run(org_date)
    set_run_font(r_m, size=Pt(9), bold=True, color=COLOR_SECONDARY, east_asia=east_asia)
    
    if desc:
        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_before = Pt(0)
        p_d.paragraph_format.space_after = Pt(2)
        p_d.paragraph_format.line_spacing = 1.15
        r_d = p_d.add_run(desc)
        set_run_font(r_d, size=Pt(9.5), color=COLOR_TEXT, east_asia=east_asia)
        
    if url:
        p_u = doc.add_paragraph()
        p_u.paragraph_format.space_before = Pt(0)
        p_u.paragraph_format.space_after = Pt(3)
        p_u.paragraph_format.keep_with_next = True
        r_ulbl = p_u.add_run(url_label)
        set_run_font(r_ulbl, size=Pt(9), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        r_u = p_u.add_run(url)
        set_run_font(r_u, size=Pt(9), color=COLOR_LINK, east_asia=east_asia)
        
    if image_filename:
        img_path = os.path.join(asset_dir, image_filename)
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.paragraph_format.space_before = Pt(2)
            p_img.paragraph_format.space_after = Pt(5)
            p_img.paragraph_format.keep_with_next = True
            r_img = p_img.add_run()
            r_img.add_picture(img_path, width=img_width)

print("Helper functions configured.")

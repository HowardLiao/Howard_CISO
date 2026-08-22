import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
import os

target_dir = "/Users/howardliao/Desktop/Howard/Howard_CISO"
photo_path = "/Users/howardliao/Desktop/Howard/Howard_CISO/assets/howard_portrait.jpg"
asset_dir = "/Users/howardliao/Desktop/Howard/Howard_CISO/assets"

COLOR_PRIMARY = RGBColor(15, 41, 66)      # #0F2942 Deep Navy
COLOR_SECONDARY = RGBColor(30, 58, 138)  # #1E3A8A Executive Blue
COLOR_SLATE = RGBColor(51, 65, 85)        # #334155 Slate
COLOR_TEXT = RGBColor(31, 41, 55)         # #1F2937 Charcoal Body Text
COLOR_MUTED = RGBColor(100, 116, 139)     # #64748B Subtitle/Meta Gray
COLOR_LINK = RGBColor(37, 99, 235)        # #2563EB Blue Link
FONT_FAMILY = 'Arial'

# Super Large Font Sizes (Another 1.5x Enlarge)
SZ_TITLE = Pt(40)         # 28pt * 1.5 = 42pt -> 40pt
SZ_H1 = Pt(28)            # 19.5pt * 1.5 = 29pt -> 28pt
SZ_H2 = Pt(24)            # 16.5pt * 1.5 = 24.75pt -> 24pt
SZ_H3 = Pt(22)            # 15.5pt * 1.5 = 23pt -> 22pt
SZ_BODY = Pt(21.5)        # 14.5pt * 1.5 = 21.75pt -> 21.5pt
SZ_BULLET = Pt(21)        # 14pt * 1.5 = 21pt
SZ_META = Pt(19.5)        # 13.5pt * 1.5 = 20.25pt -> 19.5pt

def create_base_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    return doc

def set_run_font(run, name=FONT_FAMILY, size=SZ_BODY, bold=False, italic=False, color=COLOR_TEXT, east_asia="Microsoft JhengHei"):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{east_asia}" w:cs="{name}"/>')
    rPr.append(rFonts)

def add_heading_1(doc, text, space_before=20, space_after=8, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=SZ_H1, bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="24" w:space="6" w:color="1E3A8A"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_subheading(doc, text, space_before=16, space_after=6, color=COLOR_SECONDARY, size=SZ_H2, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color, east_asia=east_asia)
    return p

def add_body_p(doc, text, space_before=0, space_after=8, line_spacing=1.25, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    set_run_font(run, size=SZ_BODY, color=COLOR_TEXT, east_asia=east_asia)
    return p

def add_bullet(doc, text, space_before=0, space_after=5, line_spacing=1.25, bold_prefix_colon=True, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.left_indent = Inches(0.35)
    
    if bold_prefix_colon and (':' in text or '：' in text) and not text.startswith('http'):
        sep = ':' if ':' in text else '：'
        parts = text.split(sep, 1)
        r1 = p.add_run(parts[0] + sep)
        set_run_font(r1, size=SZ_BULLET, bold=True, color=COLOR_TEXT, east_asia=east_asia)
        r2 = p.add_run(parts[1])
        set_run_font(r2, size=SZ_BULLET, color=COLOR_TEXT, east_asia=east_asia)
    else:
        run = p.add_run(text)
        set_run_font(run, size=SZ_BULLET, color=COLOR_TEXT, east_asia=east_asia)
    return p

def add_header(doc, name_text, title_text, subtitle_text, contact_text1, contact_text2, east_asia="Microsoft JhengHei"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_left = table.columns[0]
    col_right = table.columns[1]
    col_left.width = Inches(2.2)
    col_right.width = Inches(4.9)

    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    cell_left.width = Inches(2.2)
    cell_right.width = Inches(4.9)

    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for cell in [cell_left, cell_right]:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="none"/>'
            '<w:left w:val="none"/>'
            '<w:bottom w:val="none"/>'
            '<w:right w:val="none"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    p_photo = cell_left.paragraphs[0]
    p_photo.paragraph_format.space_before = Pt(0)
    p_photo.paragraph_format.space_after = Pt(0)
    if os.path.exists(photo_path):
        run_photo = p_photo.add_run()
        run_photo.add_picture(photo_path, width=Inches(2.0))

    p_name = cell_right.paragraphs[0]
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(4)
    run_name = p_name.add_run(name_text)
    set_run_font(run_name, size=SZ_TITLE, bold=True, color=COLOR_PRIMARY, east_asia=east_asia)

    p_title = cell_right.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(3)
    run_t1 = p_title.add_run(title_text)
    set_run_font(run_t1, size=SZ_H2, bold=True, color=COLOR_SECONDARY, east_asia=east_asia)

    p_sub = cell_right.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(6)
    run_t2 = p_sub.add_run(subtitle_text)
    set_run_font(run_t2, size=SZ_H3, bold=True, color=COLOR_SLATE, east_asia=east_asia)

    p_c1 = cell_right.add_paragraph()
    p_c1.paragraph_format.space_before = Pt(0)
    p_c1.paragraph_format.space_after = Pt(3)
    r_loc = p_c1.add_run(contact_text1)
    set_run_font(r_loc, size=SZ_META, color=COLOR_SLATE, east_asia=east_asia)

    p_c2 = cell_right.add_paragraph()
    p_c2.paragraph_format.space_before = Pt(0)
    p_c2.paragraph_format.space_after = Pt(0)
    r_c2 = p_c2.add_run(contact_text2)
    set_run_font(r_c2, size=SZ_META, color=COLOR_SLATE, east_asia=east_asia)

    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(10)
    p_sep.paragraph_format.space_after = Pt(0)

def add_exp_item(doc, role, group_note, company_line, period, desc, scopes, achs, leaving=None, add_ctx=None, east_asia="Microsoft JhengHei", labels=None):
    if labels is None:
        labels = {"scope": "Leadership Scope", "ach": "Selected Achievements & Impact", "ctx": "Additional Context", "leaving": "Reason for Leaving"}
    
    p_r = doc.add_paragraph()
    p_r.paragraph_format.space_before = Pt(18)
    p_r.paragraph_format.space_after = Pt(3)
    p_r.paragraph_format.keep_with_next = True
    r1 = p_r.add_run(role)
    set_run_font(r1, size=SZ_H2, bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    if group_note:
        p_cn = doc.add_paragraph()
        p_cn.paragraph_format.space_before = Pt(0)
        p_cn.paragraph_format.space_after = Pt(3)
        p_cn.paragraph_format.keep_with_next = True
        rcn = p_cn.add_run(group_note)
        set_run_font(rcn, size=SZ_META, color=COLOR_SLATE, east_asia=east_asia)
        
    p_m = doc.add_paragraph()
    p_m.paragraph_format.space_before = Pt(0)
    p_m.paragraph_format.space_after = Pt(4)
    p_m.paragraph_format.keep_with_next = True
    rm = p_m.add_run(f"{company_line}\n{period}")
    set_run_font(rm, size=SZ_META, bold=True, color=COLOR_SECONDARY, east_asia=east_asia)
    
    if desc:
        p_sd = doc.add_paragraph()
        p_sd.paragraph_format.space_before = Pt(4)
        p_sd.paragraph_format.space_after = Pt(6)
        p_sd.paragraph_format.keep_with_next = True
        rsd = p_sd.add_run(desc)
        set_run_font(rsd, size=SZ_BODY, italic=True, color=COLOR_TEXT, east_asia=east_asia)
        
    if scopes:
        p_sh = doc.add_paragraph()
        p_sh.paragraph_format.space_before = Pt(6)
        p_sh.paragraph_format.space_after = Pt(3)
        p_sh.paragraph_format.keep_with_next = True
        rsh = p_sh.add_run(labels["scope"])
        set_run_font(rsh, size=SZ_H3, bold=True, color=COLOR_SLATE, east_asia=east_asia)
        for s in scopes:
            add_bullet(doc, s, space_after=3.5, bold_prefix_colon=False, east_asia=east_asia)
            
    if achs:
        p_ah = doc.add_paragraph()
        p_ah.paragraph_format.space_before = Pt(6)
        p_ah.paragraph_format.space_after = Pt(3)
        p_ah.paragraph_format.keep_with_next = True
        rah = p_ah.add_run(labels["ach"])
        set_run_font(rah, size=SZ_H3, bold=True, color=COLOR_SLATE, east_asia=east_asia)
        for a in achs:
            add_bullet(doc, a, space_after=4, bold_prefix_colon=False, east_asia=east_asia)
            
    if add_ctx:
        p_ch = doc.add_paragraph()
        p_ch.paragraph_format.space_before = Pt(6)
        p_ch.paragraph_format.space_after = Pt(3)
        p_ch.paragraph_format.keep_with_next = True
        rch = p_ch.add_run(labels["ctx"])
        set_run_font(rch, size=SZ_H3, bold=True, color=COLOR_SLATE, east_asia=east_asia)
        add_body_p(doc, add_ctx, space_after=6, east_asia=east_asia)
        
    if leaving:
        p_lh = doc.add_paragraph()
        p_lh.paragraph_format.space_before = Pt(6)
        p_lh.paragraph_format.space_after = Pt(3)
        p_lh.paragraph_format.keep_with_next = True
        rlh = p_lh.add_run(labels["leaving"])
        set_run_font(rlh, size=SZ_H3, bold=True, color=COLOR_SLATE, east_asia=east_asia)
        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after = Pt(8)
        rl = p_l.add_run(leaving)
        set_run_font(rl, size=SZ_META, italic=True, color=COLOR_MUTED, east_asia=east_asia)

def add_media_entry(doc, title, org_date, desc, url=None, image_filename=None, img_width=Inches(4.2), east_asia="Microsoft JhengHei", url_label="🔗 Verified Source / DOI / URL: "):
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(14)
    p_t.paragraph_format.space_after = Pt(3)
    p_t.paragraph_format.keep_with_next = True
    r_t = p_t.add_run(title)
    set_run_font(r_t, size=SZ_H3, bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(4)
    p_meta.paragraph_format.keep_with_next = True
    r_m = p_meta.add_run(org_date)
    set_run_font(r_m, size=SZ_META, bold=True, color=COLOR_SECONDARY, east_asia=east_asia)
    
    if desc:
        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_before = Pt(0)
        p_d.paragraph_format.space_after = Pt(4)
        p_d.paragraph_format.line_spacing = 1.25
        r_d = p_d.add_run(desc)
        set_run_font(r_d, size=SZ_BODY, color=COLOR_TEXT, east_asia=east_asia)
        
    if url:
        p_u = doc.add_paragraph()
        p_u.paragraph_format.space_before = Pt(0)
        p_u.paragraph_format.space_after = Pt(6)
        p_u.paragraph_format.keep_with_next = True
        r_ulbl = p_u.add_run(url_label)
        set_run_font(r_ulbl, size=SZ_META, bold=True, color=COLOR_SLATE, east_asia=east_asia)
        r_u = p_u.add_run(url)
        set_run_font(r_u, size=SZ_META, color=COLOR_LINK, east_asia=east_asia)
        
    if image_filename:
        img_path = os.path.join(asset_dir, image_filename)
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.paragraph_format.space_before = Pt(4)
            p_img.paragraph_format.space_after = Pt(8)
            p_img.paragraph_format.keep_with_next = True
            r_img = p_img.add_run()
            r_img.add_picture(img_path, width=img_width)

# Import generator functions and update size constants
import build_trilingual_1_5x_resumes

if __name__ == "__main__":
    # We will generate ZH, EN, JA with the 1.5x enlarged fonts
    # Let's run build_trilingual_1_5x_resumes with these super large sizes
    pass

def generate_zh():
    doc = create_base_doc()
    add_header(
        doc,
        name_text="廖倫豪 博士 (Howard Liao, Ph.D.)",
        title_text="集團資安長 暨 科技副總 (Group CISO)",
        subtitle_text="全球資安治理、數位信任與架構韌性 (呈報 董事會、董事長、總經理、CEO 專用履歷)",
        contact_text1="台灣 (支援全球跨國據點)  |  行動電話：+886-975-323161  |  電子郵件：Liao.Howard@gmail.com",
        contact_text2="LinkedIn 領英：linkedin.com/in/howardliao78  |  作品集網站：https://howardliao.github.io/portfolio/",
        east_asia="Microsoft JhengHei"
    )
    
    add_heading_1(doc, "前言與高階主管職涯定位 (Executive Summary)", east_asia="Microsoft JhengHei")
    add_body_p(doc, "我是廖倫豪博士，跨國網絡科技與數位娛樂平台集團副總 暨 IT Director / 集團資安長 (Group CISO)。具備 27+ 年企業資訊科技領導力、15+ 年資訊安全治理經驗，以及 10+ 年於上市櫃、跨國與受法規監管企業擔任資安長與科技副總之高階歷練。")
    add_body_p(doc, "兼具董事會層級之資安治理高度，以及零信任 (Zero Trust)、身分存取管理 (IAM)、多雲資安 (Multi-Cloud Security)、SOC/SIEM/EDR、資安事件應變 (Incident Response)、DevSecOps、資料保護、第三方風險管理與 AI 治理 (ISO 42001) 的實戰技術深度。")
    add_body_p(doc, "精通保護關鍵業務系統、敏感機密紀錄、核心智慧財產權、企業級應用程式、雲端/SaaS 平台與跨境資料流。以務實、風險驅動的方法持續提升資安成熟度、營運韌性、稽核整備度、災難復原力與數位信任。")
    add_body_p(doc, "擁有卓越的成果紀錄，能將資安風險與技術優先級轉化為董事會與高階主管重視的商業價值、財務效益、營運指標與法規遵循結果。建構具備高擴展性的資安營運模型，在保護關鍵資產與強化業務連續性的同時，全力支援集團之跨國擴張。")
    add_body_p(doc, "過去在 IT/OT 融合與跨國多廠區工業製造環境之豐富實務，進一步奠定了駕馭複雜基礎架構、營運韌性、供應鏈整合與合作夥伴生態系安全防護的全面能力。")

    add_heading_1(doc, "董事會與高階領導價值主張 (Leadership Value Proposition)", east_asia="Microsoft JhengHei")
    add_bullet(doc, "CISO 董事會戰略治理：透過風險矩陣、RTO/RPO、SLA、ROI/TCO 與成熟度藍圖，將資安態勢、風險暴露、事件與控制落差精準轉化為具體商業衝擊與決策依據。")
    add_bullet(doc, "集團級資安營運架構：跨越企業總部、國際據點、多公有雲環境、SaaS 平台、託管服務商 (MSP) 與商業夥伴，建立全集團標準化資安營運體系。")
    add_bullet(doc, "數位信任與架構韌性：透過身分治理、安全監控、加密、完整稽核軌跡、備份、復原與即時應變機制，全方位守護關鍵系統、機密紀錄與智慧財產權。")
    add_bullet(doc, "雲原生技術深度：結合高階策略與 AWS、Azure、GCP、Kubernetes/GKE、Zero Trust、API 安全、WAF/WAAP、SIEM/SOC、EDR/XDR 與 IaC 自動化之深厚實作經驗。")
    add_bullet(doc, "AI 治理 (ISO/IEC 42001)：落地 ISO/IEC 42001 AI 管理系統標準，在推動企業負責任採用 GenAI 的同時，嚴密控管 Shadow AI、資料外洩、IP 侵權與第三方模型風險。")
    add_bullet(doc, "業務賦能與成長思維：與業務、財務、法務、人資、研發、營運、產品與 IT 團隊緊密協同，將資訊安全無縫嵌入業務增長、數位轉型與客戶信任之中。")

    add_heading_1(doc, "八大核心職能與技術治理體系 (Core Competencies)", east_asia="Microsoft JhengHei")
    comp_data_zh = [
        ("資安策略、治理與風險管理 (Cybersecurity Strategy, Governance & Risk)", [
            "集團資安策略、營運模型、政策架構、風險胃納、成熟度評估、KPI/KRI 指標設計、高階主管報告與多年期資安藍圖規劃。",
            "ISO 27001/27002、NIST CSF、Zero Trust 零信任架構、風險管理、供應商風險、資安盡職調查、稽核整備度與控制措施評估。",
            "藉由風險降低、ROI/TCO 分析、RTO/RPO、SLA、業務連續性與可量化價值實現，落實資安投資治理。",
            "董事會定期匯報、外部稽核應對、重要客戶資安審查、危機治理與高階利害關係人管理。"
        ]),
        ("數位信任、數據治理與可稽核性 (Digital Trust, Data Governance & Auditability)", [
            "關鍵業務系統、機密檔案、敏感資料、智慧財產權、企業工作流與數位平台的全面資安治理。",
            "資產盤點、業務關鍵性分級、風險評估、安全控制基準線、稽核佐證管理與標準作業程序 (SOP)。",
            "最小權限、職責分離 (SoD)、存取定期複核、日誌記錄、稽核軌跡保護、資料加密、備份復原與變更控制治理。",
            "涵蓋資料分類、存取、留存、封存、銷毀、完整性、可追溯性、可用性與可復原性之數據生命週期管理。",
            "針對 ERP、CRM、HRIS、財務系統、文件管理、資料中台、協作工具與 SaaS 生態系的安全架構與供應商審查。"
        ]),
        ("身分鑑別、零信任與資料保護 (Identity, Zero Trust & Data Protection)", [
            "企業級 IAM、SSO、MFA、條件式存取、特權存取治理 (PAM)、最小權限、員工進轉離生命週期控管與第三方身分管理。",
            "資料分級分類、DLP 防外洩、加密與金鑰管理、端點安全、安全協作、CASB/SSE 概念與跨境資料流動治理。",
            "全面防護核心智慧財產權、機密研發資料、客戶與合作夥伴資訊、財務紀錄、商業情報與內部核心資產。"
        ]),
        ("多雲、SaaS 與企業應用安全 (Cloud, SaaS & Enterprise Application Security)", [
            "AWS、Azure 與 GCP 資安治理、Landing Zones、集中化日誌、組態合規、CSPM/CNAPP 控制項、網路微隔離、WAF/WAAP、API 安全、備份與跨區災難復原。",
            "ERP/SAP、CRM、HRIS、財務軟體、資料平台、協作工具、工作流與關鍵 SaaS 服務之安全管控。",
            "雲端與 SaaS 供應商盡職調查、共同責任模型評估、架構審查、合約安全條款、持續合規保證與退場/業務連續性計畫。",
            "針對關鍵服務之跨區域 (Multi-Region) 與跨可用區 (Multi-Zone) 高可用架構設計與跨區災備機制。"
        ]),
        ("維運監控、營運韌性與事件應變 (Security Operations, Resilience & Incident Response)", [
            "SOC/MDR 營運模型、SIEM、EDR/XDR、集中式日誌聚合、偵測工程、威脅情資、弱點管理、事件應變與高階事件通報機制。",
            "勒索軟體防禦、不可變備份 (Immutable Backup)、災難復原、業務連續性、危機通訊、兵棋推演、紅藍對抗演練與復原治理。",
            "運用 ELK Stack、Graylog、Prometheus、LibreNMS、Spiceworks、雲端原生監控與系統遙測落實全域可觀測性。",
            "追蹤 MTTR、偵測覆蓋率、事件嚴重度、應變處置成效、復原能力與控制措施有效性等量化指標。"
        ]),
        ("DevSecOps、應用程式與軟體供應鏈安全 (DevSecOps, Application & Supply-Chain Security)", [
            "安全軟體開發生命週期 (SSDLC)、威脅建模、SAST、DAST、SCA、機敏密鑰管理、API 安全、安全 CI/CD、Policy-as-Code、IaC 治理、Kubernetes/GKE 資安與 SBOM。",
            "軟體供應鏈安全、開源元件合規治理、安全代碼審查、漏洞修補流程與應用程式安全規範要求。",
            "針對內部研發團隊、委外開發廠商、系統整合商、API 串接、數據管道與第三方平台之全面安全管理。"
        ]),
        ("AI 治理與安全創新 (AI Governance & Secure Innovation)", [
            "基於 ISO/IEC 42001 之 AI 治理體系、負責任 AI 政策、核准用例管理、模型與供應商風險評估、資料分級與 AI 生命週期控管。",
            "生成式 AI 資料防洩漏 (GenAI DLP)、Shadow AI 管控、安全 RAG 知識庫治理、存取控制、Prompt/資料保護、即時監控、可稽核性與人機複核 (Human-in-the-Loop)。",
            "結合 AI 輔助風險評分、異常行為偵測、威脅事件關聯分析與進階資安數據分析。"
        ]),
        ("高階領導力與跨域利害關係人管理 (Leadership & Stakeholder Management)", [
            "具備橫跨台灣、中國大陸、亞太地區、歐洲與全球託管服務夥伴的跨國跨職能領導實務。",
            "統領資安、IT 維運、雲端維運、基礎架構、SRE、DevOps、應用程式開發、服務台與外部廠商團隊。",
            "管控每年高達數百萬至千萬美元之全球 IT 與資安預算。",
            "中英文高階溝通能力；董事會定期報告；各類稽核應對；重要客戶調查評估；供應商管理；資安危機處置領導力。"
        ])
    ]
    for d_title, bullets in comp_data_zh:
        add_subheading(doc, d_title, size=SZ_H2, east_asia="Microsoft JhengHei")
        for b in bullets:
            add_bullet(doc, b, bold_prefix_colon=False, east_asia="Microsoft JhengHei")

    add_heading_1(doc, "專業歷練與職涯成就里程碑 (Professional Experience)", east_asia="Microsoft JhengHei")
    zh_labels = {"scope": "領導範疇 (Leadership Scope)", "ach": "重大成就與量化效益 (Selected Achievements & Impact)", "ctx": "補充背景 (Additional Context)", "leaving": "離職原因 (Reason for Leaving)"}
    
    # 1. 關聯集團
    add_exp_item(
        doc,
        role="集團副總 暨 資安與數位轉型負責人 (Vice President)",
        group_note="(高雄、上海、台北 / 盛欣、盛碁網絡/中國 波克/台北 芬格國際有限公司) 上市櫃公司 關聯集團",
        company_line="跨國集團 | 台灣 | 多據點運營",
        period="2025.05 – 至今",
        desc="以 CISO 級別角色主導全集團資訊安全戰略、多雲與基礎架構治理、企業營運韌性及 AI 治理，橫跨多個事業群與跨國營運據點。",
        scopes=[
            "全集團跨事業部、跨營運據點之 IT 與資安最高治理架構。",
            "帶領約 35–40 位專業人員，涵蓋資訊安全、基礎架構、應用系統、維運與外部託管服務夥伴。",
            "管控每年約 1,000–1,200 萬美元之全集團 IT 與資安總預算。"
        ],
        achs=[
            "建立基於 ISO 27001 與 NIST CSF 之集團資安治理框架，將風險管理、投資排序、韌性指標與權責機制提升至董事會層級監管。",
            "打造企業級資安風險管理體系，涵蓋資產關鍵性分級、威脅情資、弱點管理、控制措施成熟度、第三方風險與量化風險矩陣。",
            "制定緊密連結年度預算、OGSM 目標、商業優先級與可量化風險降減成果之多年期資安藍圖。",
            "針對關鍵企業應用、資料中台與夥伴協作，設計符合零信任 (Zero Trust) 原則之身分、端點、網路、雲端與存取控制架構。",
            "導入集中式 SOC/SIEM、EDR、日誌聚合、資安即時監控與應變標準作業流程，降低重大資安事件約 30%，大幅提升偵測處置效能。",
            "建立多雲資安可觀測性與組態合規儀表板，提升高階主管能見度、稽核佐證效率、重要客戶盡職調查回應與弱點修復治理。",
            "推動導入 ISO/IEC 42001 AI 管理系統，訂定 AI 使用規範、風險評估、資料防護標準、模型生命週期控制與第三方 AI 服務治理。",
            "透過 ROI/TCO、RTO/RPO、SLA、風險降減幅度與稽核整備度分析，成功爭取零信任、SOC/SIEM、災難復原與雲端資安等重大投資贊助。",
            "有效協調 IT、財務、法務、人資、研發、業務營運與外部服務供應商，推進資安與數位轉型專案。"
        ],
        leaving="尋求位於台灣北部、能與長期家庭生活規劃、職涯發展及集團級資安長治理目標高度契合的高階主管機會。",
        labels=zh_labels
    )

    # 2. 隆中網絡
    add_exp_item(
        doc,
        role="資訊處長 / IT Director",
        group_note="隆中網絡股份有限公司 / GameSparcs (總部在台中，並在台北、洛杉磯、雪梨、馬爾他、杭州與成都設有據點 / 向上國際XSGames、隆中網絡、萬國遊戲、海淯遊戲、VIVIDGAMING、銀河網絡、浩天遊戲、晶綺科技) 上市櫃公司",
        company_line="上市櫃全球線上娛樂與遊戲平台營運商 | 台中, 台灣",
        period="2022.09 – 2025.04",
        desc="主導全球雲端架構、資安治理、數位平台韌性、DevOps 效能與高可用性維運，支撐服務全球數百萬玩家之上市櫃線上平台。",
        scopes=[
            "帶領約 30–35 位專業技術人員，涵蓋雲端維運、DevOps、SRE、資訊安全、基礎架構與應用系統支援。",
            "負責每年約 1,200–1,400 萬美元之多雲基礎架構與資安預算。",
            "全面支撐需 24/7 全天候高可用、巨量交易處理與跨區域高韌性之高並發大型面向客戶數位平台。"
        ],
        achs=[
            "設計並交付以 GKE 為核心之多可用區、多區域架構與跨區災備機制，確保 24/7 服務不中斷，創下關鍵服務 100% 零停機 (Zero Outage) 紀錄。",
            "建立集中式可觀測性、APM、SIEM、雲端日誌聚合、即時監控與應變機制，透過即時事件關聯分析將 MTTR 縮短約 30%。",
            "落實符合 ISO 27001 與 NIST CSF 之資安政策與控制措施，涵蓋 IAM、資產管理、安全開發、雲端維運、事件應變、供應商治理與稽核軌跡。",
            "透過 IAM、SSO、MFA、最小權限與安全遠端存取導入零信任原則，強化分散式團隊、特權用戶與外部合作夥伴之防護力。",
            "強化外部攻擊面防禦，導入 API 安全防護、WAF/WAAP、DDoS 緩解、防爬蟲機制、速率限制、API 鑑權與網路安全治理。",
            "建立資安應變 Playbooks、應變演練、紅藍對抗推演、危機通訊程序與復原治理，大幅提升組織資安整備度並降低資安事件衝擊。",
            "將 SAST、DAST、機敏密鑰管理、CI/CD 安全控制、IaC 治理、容器安全與軟體供應鏈安全無縫整合至工程研發與部署流程。",
            "推動多雲 FinOps 雲端財務營運實踐，在業務流量劇增期間維持極致效能與穩定性的同時，達成約 30% 雲端成本優化。",
            "憑藉兼具超高可用性、營運韌性與成本最佳化之 GKE 多雲架構，榮獲「雲端架構卓越獎 (Cloud Architecture Excellence Award)」。"
        ],
        leaving="尋求具備更廣泛全集團治理權責、能與長期職涯目標更緊密契合的資深資安長領導職位。",
        labels=zh_labels
    )

    # 3. 泓晏科技
    add_exp_item(
        doc,
        role="資訊處長 / IT Director",
        group_note="泓晏科技 (總部在深圳，並在新竹、杭州與成都設有據點 / 泓晏科技)",
        company_line="電子製造與高科技技術服務集團 | 台灣 與 中國大陸",
        period="2018.03 – 2022.08",
        desc="統領企業 IT、資訊安全、數據治理、數位轉型及跨越台灣與中國大陸之多廠區多據點科技營運體系。",
        scopes=[
            "管理約 30–40 位專業人員，橫跨 IT、資安、企業應用、基礎架構、維運與技術服務範疇。",
            "管控每年約 500–700 萬美元之 IT 與資訊安全年度預算。",
            "全面支援多廠區營運、研發中心、供應鏈、工程、財務、採購與外部合作夥伴環境。"
        ],
        achs=[
            "為企業與製造營運環境建立符合 ISO 27001 之資安治理、風險評估、資產盤點、身分與存取控制、網路微隔離、弱點管理與可稽核性。",
            "為 ERP、PLM、供應鏈、工作流、數據分析與企業核心系統設計安全架構，提升高可用性、可追溯性、資料防護與業務連續性。",
            "主導 ERP、MES、PLM、APS、HR 與 BI 分析平台之安全整合，實現跨業務流程之受控資料流、身分鑑別、傳輸加密、稽核日誌與管理能見度。",
            "建置數據治理與商業智慧 (BI) 能力，有效支援庫存智慧管理、異常分析、營運決策與併購後數據資產整合。",
            "建立兼顧資安防護與生產營運連續性之風險驅動漏洞修補、修補管理、事件處置、備份與業務連續性程序。",
            "橫跨台灣與大陸協調 IT、工程、製造營運、採購、財務及外部服務供應商之資安與營運韌性專案。",
            "於集團被中國大型製造業集團收購期間，主導組織 IT 整合與數據資產之盡職調查。"
        ],
        add_ctx="深厚的 IT/OT 融合經驗，為複雜基礎架構、供應鏈整合、夥伴生態系與跨國製造運營提供堅實支撐。",
        leaving="隨公司順利併入大型製造集團後，為追求具備更廣闊戰略發展之資訊安全與科技高階領導機會而離任。",
        labels=zh_labels
    )

    # 4. 凌網科技
    add_exp_item(
        doc,
        role="資訊部經理 / IT Manager",
        group_note="凌網科技股份有限公司 (總部在新竹，並在台北、台中、高雄、泰國、北京設有據點 / 凌網科技、凌網知識) 上市櫃公司",
        company_line="上市櫃大型資訊服務集團 | 台灣、中國大陸 與 泰國",
        period="2014.05 – 2018.02",
        desc="負責亞太區 IT 維運、服務管理 (ITSM)、資料中心戰略、基礎架構治理與跨越台北、新竹、台中、北京與泰國之跨境技術支援。",
        scopes=[
            "管理 34–45 位專業工程師，橫跨基礎架構、企業應用、Service Desk、技術支援與機房維運。",
            "主導區域 IT 預算、資料中心營運、備份架構、災難復原計畫與服務管理流程優化。"
        ],
        achs=[
            "導入 ITSM 與跨區域服務管理機制，提升亞太各分公司之服務一致性、當責性、事件處置速度與營運透明度。",
            "制定區域資料中心戰略、高可用架構、備份治理與災難復原方案，強力支撐關鍵業務服務。",
            "引入 PMP 與敏捷交付實務，強化跨國專案治理、交付可預測性與利害關係人協作效能。",
            "為跨國地理分散據點建立一致化之基礎架構標準、資安營運規範與技術服務體系。"
        ],
        labels=zh_labels
    )

    # 5. 光南集團
    add_exp_item(
        doc,
        role="資訊部經理 / 數位轉型負責人",
        group_note="光南集團 / 峰晨集團 – 艾居電腦 (總部在台中，並在台北、深圳、上海、張家港設有據點 / 光南集團) 上市櫃公司",
        company_line="零售通路、供應鏈物流與科技服務集團 | 台灣 與 中國大陸",
        period="2011.07 – 2014.04",
        desc="於零售與流通集團併購期間，主導企業核心系統整合、商業數據治理與數位轉型推動。",
        scopes=[
            "管理 15–20 位專業 IT 技術團隊成員。",
            "負責 ERP、POS、CRM、企業數據整合、業務連續性、應用系統治理與轉型預算。"
        ],
        achs=[
            "整併 ERP、POS、CRM 與數據分析中台，全面提升集團對門市銷售、庫存周轉、會員行為與營運績效之全局掌控力。",
            "主導跨國併購期間之數據清洗、主數據治理 (MDM)、系統遷移與業務系統整合，確保業務不中斷。",
            "建構整合式企業數據架構，全面賦能管理階層之即時報表分析與策略決策支援。"
        ],
        labels=zh_labels
    )

    # 6. Borland
    add_exp_item(
        doc,
        role="資訊部經理 / IT Manager",
        group_note="美商寶藍 (Borland Taiwan) 跨國外商公司",
        company_line="跨國軟體與生命週期管理外商集團 | 台灣 與 中國大陸",
        period="2008.03 – 2011.06",
        desc="為台灣與中國大陸研發及營運團隊建構跨國 IT 協作平台、高可用企業系統、基礎架構服務與數據保護機制。",
        scopes=None,
        achs=[
            "為企業級應用程式與關鍵交易環境設計高可用資料庫、應用中台、備份與災難復原架構。",
            "支援軟體生命週期管理 (ALM) 平台、工程協作系統與區域 IT 服務之全面部署與維運。",
            "維持跨國運營環境中之高度系統穩定性、資料安全保護與跨境研發協作效能。"
        ],
        labels=zh_labels
    )

    # 7. Sybase
    add_exp_item(
        doc,
        role="資訊部經理 / IT Manager",
        group_note="美商賽貝斯 (Sybase) 跨國外商公司",
        company_line="跨國頂級資料庫與企業軟體外商集團 | 亞太區",
        period="2002.08 – 2008.02",
        desc="主導亞太區資料中心整併、基礎架構標準化、高可用資料庫架構、區域 IT 維運，並為電信與金融級客戶提供關鍵任務技術支援。",
        scopes=[
            "管理亞太區基礎架構團隊約 10–15 位資深工程師。",
            "全面支援企業級大型資料庫、高可用叢集、容錯移轉、備份與資料中心營運。"
        ],
        achs=[
            "於亞太區據點建置統一 Active Directory 目錄服務、資料中心整併、高可用資料庫與自動容錯移轉策略。",
            "為電信運營商與金融機構關鍵任務環境提供電信級 (Carrier-Grade) 資料庫調優、高可用性、災難復原與效能解決方案。",
            "在跨國區域整併與全球大客戶專案執行期間，確保 100% 業務營運連續性與各項稽核法規遵循。"
        ],
        labels=zh_labels
    )

    add_heading_1(doc, "專業國際證照資格 (Certifications)", east_asia="Microsoft JhengHei")
    for cert in [
        "ISO 27001 資訊安全管理系統 — 主任稽核員 / 內部稽核員 (Lead Auditor)",
        "ISO/IEC 42001 人工智慧管理系統 (AIMS) — AI 治理主任稽核員認證",
        "國際專案管理師認證 — PMP (Project Management Professional)",
        "敏捷專案大師認證 — CSM (Certified ScrumMaster)",
        "Oracle 官方認證專業專家 — OCP (Oracle Certified Professional)",
        "Oracle 官方認證專員 — OCA (Oracle Certified Associate)",
        "ESG 永續規劃師證照 (ESG Sustainability Planner)"
    ]:
        add_bullet(doc, cert, bold_prefix_colon=False, east_asia="Microsoft JhengHei")

    add_heading_1(doc, "正規學術學位 (Academic Education)", east_asia="Microsoft JhengHei")
    edu_list_zh = [
        ("資訊科技管理 博士 (Ph.D. in IT Management)", "朝陽科技大學 資訊科技管理研究所 (College of Informatics)\n2009.09 – 2013.06", "研究領域：資訊安全治理、雲端安全、AI 治理、數據治理、軟體工程、DevSecOps 與企業數位轉型戰略。"),
        ("資訊科學 碩士 (Master of Science in Information Science)", "國立中興大學 資訊科學與工程研究所\n2004.09 – 2006.06", "專注領域：現代密碼學 (Modern Cryptography)、資料探勘、數據治理、軟體工程、系統分析與企業資訊系統架構。"),
        ("資訊科學與工程 學士 (Bachelor of Computer Science)", "國立中興大學 資訊科學與工程學系\n2000.09 – 2004.06", "專注領域：軟體開發生命週期 (SDLC)、軟體工程、系統分析、企業架構與應用程式開發。")
    ]
    for d, s, desc in edu_list_zh:
        add_subheading(doc, d, size=SZ_H3, east_asia="Microsoft JhengHei")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(s)
        set_run_font(r, size=SZ_META, bold=True, color=COLOR_SECONDARY, east_asia="Microsoft JhengHei")
        add_body_p(doc, desc, space_after=6, east_asia="Microsoft JhengHei")

    add_heading_1(doc, "獲獎榮譽、大會演講、國際期刊論文與媒體報導", east_asia="Microsoft JhengHei")
    add_subheading(doc, "1. Keynote 大會主題演講與實務分享", size=SZ_H2, east_asia="Microsoft JhengHei")
    add_media_entry(doc, "MongoDB.local Taipei Keynote 大會主講《運用 MongoDB Atlas 建構全球規模遊戲資料管理平台》", "MongoDB 原廠大會 (Taipei) | 2024.09", "擔任 Keynote 主講人，公開發表全球跨區高可用資料架構，闡述如何藉由 MongoDB Atlas 全託管雲端資料庫兼具高擴展與 99.995% 高可用性，支撐全球百萬玩家跨區即時對弈與高併發讀寫需求。", url="https://www.ithome.com.tw/pr/163534", image_filename="2024_MongoDB02.jpg", img_width=Inches(4.2), east_asia="Microsoft JhengHei", url_label="🔗 查證佐證出處 / URL：")
    add_media_entry(doc, "CIO Taiwan 價值學院第十七屆大會 Keynote 演講《洞見機遇-雲端管理與實務》", "CIO Taiwan 價值學院 (Taipei) | 2024.05", "受邀擔任 Keynote 主講人，向與會上市櫃企業 CIO 與 CISO 分享多雲資源治理、安全合規與 FinOps 實務落地心法，探討雲端現代化治理與架構韌性。", url="https://www.cio.com.tw/events/value-index-2/", image_filename="問題比答案重要.png", img_width=Inches(3.8), east_asia="Microsoft JhengHei", url_label="🔗 查證佐證出處 / URL：")
    add_media_entry(doc, "行政院人事行政總處 e等公務園+學習平臺《洞見機遇-雲端管理與實務》大師講座數位課程", "行政院人事行政總處 e等公務園數位學習平台 | 2024.06", "受邀為行政院公務同仁錄製《洞見機遇-雲端管理與實務》大師講座數位課程，推廣公部門雲端轉型、資安韌性與現代化資訊治理思維，列入國家公務員數位學習專屬教材。", url="https://elearn.hrd.gov.tw/info/10042804", image_filename="e等公務園__洞見機遇-雲端管理與實務01.png", img_width=Inches(3.8), east_asia="Microsoft JhengHei", url_label="🔗 查證佐證出處 / URL：")
    add_media_entry(doc, "Google Cloud 官方全球客戶成功案例影音專訪 (GameSparcs APAC Customer Success Story)", "Google Cloud APAC 官方專訪 (YouTube) | 2024.03", "接受 Google Cloud 官方採訪，深入解析如何運用 GKE 多雲架構與自動擴縮容技術，支撐全球百萬級高併發手遊連線營運，創下 100% Zero Outage 零停機紀錄。", url="https://youtu.be/_kTZSZ_0lNE?si=CT2lo8c4IF0zI1Ki", east_asia="Microsoft JhengHei", url_label="🔗 官方影片連結：")

    add_subheading(doc, "2. 國際期刊論文與國家學術典藏", size=SZ_H2, east_asia="Microsoft JhengHei")
    add_media_entry(doc, "Springer SCI 國際頂級期刊論文：Reversible secret-image sharing with high visual quality", "Multimedia Tools and Applications (Springer Nature, Vol. 74, Pages 10603–10626) | 2014.06", "共同作者：Ching-Chiuan Lin, Lun-Hao Liao (廖倫豪), Kuo-Feng Hwang, Shih-Chieh Chen。本論文提出基於高視覺品質陰影圖像之可逆秘密影像共享技術，在密碼學與資安影像傳輸領域具高度學術影響力。", url="https://link.springer.com/article/10.1007/s11042-012-1190-1", east_asia="Microsoft JhengHei", url_label="🔗 DOI 永久出版品連結：")
    add_media_entry(doc, "國際學術期刊論文：Utilizing GIS and GPS in Designing a Trilingual Tourist APP", "Applied Science and Management Research (Vol. 2) | 2015.05", "探討結合地理資訊系統 (GIS) 與全球定位系統 (GPS) 開發三語觀光導覽系統之行動應用架構設計與實作。", east_asia="Microsoft JhengHei")
    add_media_entry(doc, "國家圖書館博碩士學位論文法定典藏《整合系統的商業自助式入口網站植基於調適性服務導向架構資訊科技治理之研究》", "國家圖書館臺灣博碩士論文知識加值系統 (典藏代碼: 106IKTC0183002) | 2014.10", "博士論文研究：結合 ITIL V3 與 SOA 服務導向架構建立企業級自助式入口網站 IT 治理模型，並於國家圖書館永久法定典藏，擔任學術研究指導學者典藏代碼 106IKTC0183002。", url="https://ndltd.ncl.edu.tw/", east_asia="Microsoft JhengHei", url_label="🔗 國圖學術系統：")

    add_subheading(doc, "3. 主流科技媒體實名專訪與專題報導", size=SZ_H2, east_asia="Microsoft JhengHei")
    add_media_entry(doc, "CIO Taiwan 官方專訪《【專訪】隆中網絡 GameSparcs IT Director Howard Liao | 善用公有雲服務 搶攻全球遊戲商機》", "CIO Taiwan 雜誌 (採訪／施鑫澤‧文／林裕洋‧刊期／2024.05) | 2024.05", "CIO Taiwan 總編輯專題專訪，實名刊登 Howard Liao, PhD 廖博士之多雲架構佈局，深入剖析採用多雲網路 (MCN) 與 Kubernetes 原生架構橫跨公有雲與邊緣運算，強化全球發行競爭力與 FinOps 降本 30%。", url="https://www.cio.com.tw/interview-howard-liao-online-gamesparcs-it-director/", image_filename="2024_CIO報導.png", img_width=Inches(3.8), east_asia="Microsoft JhengHei", url_label="🔗 媒體報導連結：")
    add_media_entry(doc, "iThome 電腦報專題企劃《【iThome 專題企劃】隆中網絡運用 MongoDB Atlas 建構遊戲資料管理平台，打造全方位休閒娛樂平台》", "iThome 電腦報官方專題企劃 | 2024.09", "技術專題企劃實名專訪廖倫豪博士，闡述以 MongoDB Atlas 建置多區域即時數據中台，兼具 99.995% 高可用性與高擴展性，成功支撐全球跨區高併發讀寫需求。", url="https://www.ithome.com.tw/pr/163534", image_filename="IThome_MangoDB.png", img_width=Inches(4.0), east_asia="Microsoft JhengHei", url_label="🔗 媒體報導連結：")
    add_media_entry(doc, "DIGITIMES 科技網報導：MongoDB 8.0 問世，助企業接軌生成式 AI (隆中網絡等企業案例)", "DIGITIMES 科技網 | 2024.09", "DIGITIMES 深度報導 MongoDB 8.0 發表會，實名收錄隆中網絡技術團隊架構分享與 AI 數據治理之成功應用案例。", url="https://www.digitimes.com.tw/tech/dt/n/shwnws.asp?id=0000704377_KI14VVDK6CQ8TB8OXC3GJ", east_asia="Microsoft JhengHei", url_label="🔗 媒體報導連結：")

    add_subheading(doc, "4. 業界獲獎榮譽與社會公益責任 (CSR)", size=SZ_H2, east_asia="Microsoft JhengHei")
    add_bullet(doc, "雲端架構卓越獎 (Cloud Architecture Excellence Award) — 表彰 GKE 多雲架構、高可用性、營運韌性與 FinOps 降本 30% 之卓越成效。", bold_prefix_colon=False, east_asia="Microsoft JhengHei")
    add_bullet(doc, "受邀擔任資安高峰會 (Cybersecurity Summit) Keynote 主講人，分享混合雲資安、零信任架構、ISO 27001、ISO/IEC 42001 與營運韌性實務。", bold_prefix_colon=False, east_asia="Microsoft JhengHei")
    add_bullet(doc, "受邀於 CIO 價值學院及技術社群分享多雲治理、AI 治理、企業數位轉型與資安戰略。", bold_prefix_colon=False, east_asia="Microsoft JhengHei")
    add_media_entry(doc, "財團法人康善基金會 (Cancer NoNo Foundation) 董事 兼 資訊長 (CIO)", "財團法人康善基金會 (CSR 永續公益) | 2022 – 至今", "擔任基金會董事兼資訊長，推動兒童健康飲食教育、便當童話劇、數位公益與醫療資訊安全系統整合，落實企業社會責任 (CSR) 與數位治理。", url="https://www.cancer-nono.org.tw/web/about/page.php?lang=zh_tw&scid=7&sid=5", east_asia="Microsoft JhengHei", url_label="🔗 基金會官方連結：")

    out_file = os.path.join(target_dir, "Howard_Liao_CISO_Resume_ZH.docx")
    out_default = os.path.join(target_dir, "Howard_Liao_CISO_Resume.docx")
    doc.save(out_file)
    doc.save(out_default)
    print(f"Generated Super Large ZH: {out_file}")

if __name__ == "__main__":
    generate_zh()
    import build_trilingual_1_5x_resumes
    build_trilingual_1_5x_resumes.SZ_TITLE = SZ_TITLE
    build_trilingual_1_5x_resumes.SZ_H1 = SZ_H1
    build_trilingual_1_5x_resumes.SZ_H2 = SZ_H2
    build_trilingual_1_5x_resumes.SZ_H3 = SZ_H3
    build_trilingual_1_5x_resumes.SZ_BODY = SZ_BODY
    build_trilingual_1_5x_resumes.SZ_BULLET = SZ_BULLET
    build_trilingual_1_5x_resumes.SZ_META = SZ_META
    build_trilingual_1_5x_resumes.generate_en()
    build_trilingual_1_5x_resumes.generate_ja()
    print("All 3 resumes regenerated with super large font sizes!")

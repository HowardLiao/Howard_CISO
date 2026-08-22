import os
import json
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
import sys

target_dir = "/Users/howardliao/Desktop/Howard/Howard_CISO"
photo_path = "/Users/howardliao/Library/Application Support/Hermes/composer-images/composer_2026-08-21_07-55-16-061_53c1dd.jpg"
asset_dir = "/Users/howardliao/Desktop/Howard/Howard_CISO/assets"

COLOR_PRIMARY = RGBColor(15, 41, 66)      # #0F2942 Deep Navy
COLOR_SECONDARY = RGBColor(30, 58, 138)  # #1E3A8A Executive Blue
COLOR_SLATE = RGBColor(51, 65, 85)        # #334155 Slate
COLOR_TEXT = RGBColor(31, 41, 55)         # #1F2937 Charcoal Body Text
COLOR_MUTED = RGBColor(100, 116, 139)     # #64748B Subtitle/Meta Gray
COLOR_LINK = RGBColor(37, 99, 235)        # #2563EB Blue Link
FONT_FAMILY = 'Arial'

def create_base_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    return doc

def set_run_font(run, name=FONT_FAMILY, size=Pt(10), bold=False, italic=False, color=COLOR_TEXT, east_asia="Microsoft JhengHei"):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{east_asia}" w:cs="{name}"/>')
    rPr.append(rFonts)

def add_heading_1(doc, text, space_before=14, space_after=4, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=Pt(13), bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="12" w:space="3" w:color="1E3A8A"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_subheading(doc, text, space_before=8, space_after=2, color=COLOR_SECONDARY, size=Pt(10.5), east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color, east_asia=east_asia)
    return p

def add_body_p(doc, text, space_before=0, space_after=4, line_spacing=1.15, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    set_run_font(run, size=Pt(9.5), color=COLOR_TEXT, east_asia=east_asia)
    return p

def add_bullet(doc, text, space_before=0, space_after=2, line_spacing=1.15, bold_prefix_colon=True, east_asia="Microsoft JhengHei"):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.left_indent = Inches(0.22)
    
    if bold_prefix_colon and (':' in text or '：' in text) and not text.startswith('http'):
        sep = ':' if ':' in text else '：'
        parts = text.split(sep, 1)
        r1 = p.add_run(parts[0] + sep)
        set_run_font(r1, size=Pt(9.5), bold=True, color=COLOR_TEXT, east_asia=east_asia)
        r2 = p.add_run(parts[1])
        set_run_font(r2, size=Pt(9.5), color=COLOR_TEXT, east_asia=east_asia)
    else:
        run = p.add_run(text)
        set_run_font(run, size=Pt(9.5), color=COLOR_TEXT, east_asia=east_asia)
    return p

def add_header(doc, name_text, title_text, subtitle_text, contact_text1, contact_text2, east_asia="Microsoft JhengHei"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_left = table.columns[0]
    col_right = table.columns[1]
    col_left.width = Inches(1.5)
    col_right.width = Inches(5.5)

    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    cell_left.width = Inches(1.5)
    cell_right.width = Inches(5.5)

    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for cell in [cell_left, cell_right]:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="none"/>'
            '<w:left w:val="none"/>'
            '<w:bottom w:val="none"/>'
            '<w:right w:val="none"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    p_photo = cell_left.paragraphs[0]
    p_photo.paragraph_format.space_before = Pt(0)
    p_photo.paragraph_format.space_after = Pt(0)
    if os.path.exists(photo_path):
        run_photo = p_photo.add_run()
        run_photo.add_picture(photo_path, width=Inches(1.35))

    p_name = cell_right.paragraphs[0]
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run(name_text)
    set_run_font(run_name, size=Pt(19), bold=True, color=COLOR_PRIMARY, east_asia=east_asia)

    p_title = cell_right.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(1)
    run_t1 = p_title.add_run(title_text)
    set_run_font(run_t1, size=Pt(11), bold=True, color=COLOR_SECONDARY, east_asia=east_asia)

    p_sub = cell_right.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(4)
    run_t2 = p_sub.add_run(subtitle_text)
    set_run_font(run_t2, size=Pt(10), bold=True, color=COLOR_SLATE, east_asia=east_asia)

    p_c1 = cell_right.add_paragraph()
    p_c1.paragraph_format.space_before = Pt(0)
    p_c1.paragraph_format.space_after = Pt(1)
    r_loc = p_c1.add_run(contact_text1)
    set_run_font(r_loc, size=Pt(9), color=COLOR_SLATE, east_asia=east_asia)

    p_c2 = cell_right.add_paragraph()
    p_c2.paragraph_format.space_before = Pt(0)
    p_c2.paragraph_format.space_after = Pt(0)
    r_c2 = p_c2.add_run(contact_text2)
    set_run_font(r_c2, size=Pt(9), color=COLOR_SLATE, east_asia=east_asia)

    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after = Pt(0)

def add_exp_item(doc, role, group_note, company_line, period, desc, scopes, achs, leaving=None, add_ctx=None, east_asia="Microsoft JhengHei", labels=None):
    if labels is None:
        labels = {"scope": "Leadership Scope", "ach": "Selected Achievements & Impact", "ctx": "Additional Context", "leaving": "Reason for Leaving"}
    
    p_r = doc.add_paragraph()
    p_r.paragraph_format.space_before = Pt(10)
    p_r.paragraph_format.space_after = Pt(1)
    p_r.paragraph_format.keep_with_next = True
    r1 = p_r.add_run(role)
    set_run_font(r1, size=Pt(11.5), bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    if group_note:
        p_cn = doc.add_paragraph()
        p_cn.paragraph_format.space_before = Pt(0)
        p_cn.paragraph_format.space_after = Pt(1)
        p_cn.paragraph_format.keep_with_next = True
        rcn = p_cn.add_run(group_note)
        set_run_font(rcn, size=Pt(9.5), color=COLOR_SLATE, east_asia=east_asia)
        
    p_m = doc.add_paragraph()
    p_m.paragraph_format.space_before = Pt(0)
    p_m.paragraph_format.space_after = Pt(2)
    p_m.paragraph_format.keep_with_next = True
    rm = p_m.add_run(f"{company_line}\n{period}")
    set_run_font(rm, size=Pt(9.5), bold=True, color=COLOR_SECONDARY, east_asia=east_asia)
    
    if desc:
        p_sd = doc.add_paragraph()
        p_sd.paragraph_format.space_before = Pt(2)
        p_sd.paragraph_format.space_after = Pt(3)
        p_sd.paragraph_format.keep_with_next = True
        rsd = p_sd.add_run(desc)
        set_run_font(rsd, size=Pt(9.5), italic=True, color=COLOR_TEXT, east_asia=east_asia)
        
    if scopes:
        p_sh = doc.add_paragraph()
        p_sh.paragraph_format.space_before = Pt(3)
        p_sh.paragraph_format.space_after = Pt(1)
        p_sh.paragraph_format.keep_with_next = True
        rsh = p_sh.add_run(labels["scope"])
        set_run_font(rsh, size=Pt(9.5), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        for s in scopes:
            add_bullet(doc, s, space_after=1.5, bold_prefix_colon=False, east_asia=east_asia)
            
    if achs:
        p_ah = doc.add_paragraph()
        p_ah.paragraph_format.space_before = Pt(3)
        p_ah.paragraph_format.space_after = Pt(1)
        p_ah.paragraph_format.keep_with_next = True
        rah = p_ah.add_run(labels["ach"])
        set_run_font(rah, size=Pt(9.5), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        for a in achs:
            add_bullet(doc, a, space_after=2, bold_prefix_colon=False, east_asia=east_asia)
            
    if add_ctx:
        p_ch = doc.add_paragraph()
        p_ch.paragraph_format.space_before = Pt(3)
        p_ch.paragraph_format.space_after = Pt(1)
        p_ch.paragraph_format.keep_with_next = True
        rch = p_ch.add_run(labels["ctx"])
        set_run_font(rch, size=Pt(9.5), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        add_body_p(doc, add_ctx, space_after=3, east_asia=east_asia)
        
    if leaving:
        p_lh = doc.add_paragraph()
        p_lh.paragraph_format.space_before = Pt(3)
        p_lh.paragraph_format.space_after = Pt(1)
        p_lh.paragraph_format.keep_with_next = True
        rlh = p_lh.add_run(labels["leaving"])
        set_run_font(rlh, size=Pt(9.5), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after = Pt(4)
        rl = p_l.add_run(leaving)
        set_run_font(rl, size=Pt(9), italic=True, color=COLOR_MUTED, east_asia=east_asia)

def add_media_entry(doc, title, org_date, desc, url=None, image_filename=None, img_width=Inches(3.2), east_asia="Microsoft JhengHei", url_label="🔗 Verified Source / DOI / URL: "):
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(7)
    p_t.paragraph_format.space_after = Pt(1)
    p_t.paragraph_format.keep_with_next = True
    r_t = p_t.add_run(title)
    set_run_font(r_t, size=Pt(10.5), bold=True, color=COLOR_PRIMARY, east_asia=east_asia)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(2)
    p_meta.paragraph_format.keep_with_next = True
    r_m = p_meta.add_run(org_date)
    set_run_font(r_m, size=Pt(9), bold=True, color=COLOR_SECONDARY, east_asia=east_asia)
    
    if desc:
        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_before = Pt(0)
        p_d.paragraph_format.space_after = Pt(2)
        p_d.paragraph_format.line_spacing = 1.15
        r_d = p_d.add_run(desc)
        set_run_font(r_d, size=Pt(9.5), color=COLOR_TEXT, east_asia=east_asia)
        
    if url:
        p_u = doc.add_paragraph()
        p_u.paragraph_format.space_before = Pt(0)
        p_u.paragraph_format.space_after = Pt(3)
        p_u.paragraph_format.keep_with_next = True
        r_ulbl = p_u.add_run(url_label)
        set_run_font(r_ulbl, size=Pt(9), bold=True, color=COLOR_SLATE, east_asia=east_asia)
        r_u = p_u.add_run(url)
        set_run_font(r_u, size=Pt(9), color=COLOR_LINK, east_asia=east_asia)
        
    if image_filename:
        img_path = os.path.join(asset_dir, image_filename)
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.paragraph_format.space_before = Pt(2)
            p_img.paragraph_format.space_after = Pt(5)
            p_img.paragraph_format.keep_with_next = True
            r_img = p_img.add_run()
            r_img.add_picture(img_path, width=img_width)

# Load DATA from build_index_html.py
with open("/Users/howardliao/Desktop/Howard/Howard_CISO/build_index_html.py", "r", encoding="utf-8") as f:
    orig = f.read()

data_idx = orig.find("const DATA = {")
end_idx = orig.find("    let currentLang = 'en';")
raw_data_block = orig[data_idx:end_idx].strip()

# We can parse the DATA dictionary via python
# Or construct clean generators for EN, ZH, JA

# 1. GENERATE ENGLISH (Howard_Liao_CISO_Resume_EN.docx)
def build_en_docx():
    doc = create_base_doc()
    add_header(
        doc,
        name_text="Howard Liao, Ph.D.",
        title_text="Group Chief Information Security Officer (CISO)",
        subtitle_text="Global Cybersecurity, Digital Trust & Resilience",
        contact_text1="Taiwan (Open to Global Sites)  |  Mobile: +886-975-323161  |  Email: Liao.Howard@gmail.com",
        contact_text2="LinkedIn: linkedin.com/in/howardliao78  |  Portfolio: https://howardliao.github.io/portfolio/"
    )
    
    add_heading_1(doc, "Executive Profile")
    add_body_p(doc, "CISO-level technology executive with 27+ years of enterprise IT leadership, 15+ years of cybersecurity experience, and 10+ years leading cybersecurity strategy, cloud governance, digital resilience, and enterprise transformation across publicly listed, multinational, and regulated business environments.")
    add_body_p(doc, "Combines board-level cybersecurity leadership with hands-on technical depth in Zero Trust, identity and access management, multi-cloud security, SOC/SIEM/EDR, incident response, DevSecOps, data protection, third-party risk management, and AI governance.")
    add_body_p(doc, "Experienced in securing business-critical systems, sensitive records, intellectual property, enterprise applications, cloud/SaaS platforms, and cross-border data flows. Applies a pragmatic, risk-based approach to improving security maturity, operational resilience, audit readiness, recoverability, and digital trust.")
    add_body_p(doc, "Proven record of translating cybersecurity risks and technical priorities into measurable business, financial, operational, and compliance outcomes for boards and executive leadership. Builds scalable security operating models that enable international growth while protecting critical assets and improving business continuity.")
    add_body_p(doc, "Prior experience in IT/OT convergence and multi-site industrial environments provides additional capability for complex infrastructure, operational resilience, supply-chain integration, and partner ecosystem security.")

    add_heading_1(doc, "Leadership Value Proposition")
    add_bullet(doc, "CISO and board leadership: Translates security posture, risk exposure, incidents, investment priorities, and control gaps into business impact through risk matrices, RTO/RPO, SLA, ROI/TCO, and maturity roadmaps.")
    add_bullet(doc, "Group security governance: Establishes group-wide security operating models across headquarters, international sites, cloud environments, SaaS platforms, managed-service providers, and business partners.")
    add_bullet(doc, "Digital trust and resilience: Protects critical enterprise systems, confidential records, intellectual property, and sensitive information through identity governance, security monitoring, encryption, auditability, backup, recovery, and incident readiness.")
    add_bullet(doc, "Cloud-native technical depth: Combines executive strategy with practical experience across AWS, Azure, GCP, Kubernetes/GKE, Zero Trust, API security, WAF/WAAP, SIEM/SOC, EDR/XDR, DevSecOps, and infrastructure automation.")
    add_bullet(doc, "AI governance: Implements ISO/IEC 42001-based AI governance to enable responsible AI adoption while controlling shadow AI, data leakage, IP exposure, third-party model risk, and operational risk.")
    add_bullet(doc, "Business-enabling mindset: Works closely with business, finance, legal, HR, engineering, operations, product, and IT leaders to embed cybersecurity into growth, transformation, and customer trust.")

    add_heading_1(doc, "Core Competencies")
    comp_data = [
        ("Cybersecurity Strategy, Governance & Risk", [
            "Group cybersecurity strategy, operating model, policy architecture, risk appetite, maturity assessments, KPI/KRI design, executive reporting, and multi-year security roadmaps.",
            "ISO 27001/27002, NIST Cybersecurity Framework, Zero Trust architecture, risk management, vendor risk, security due diligence, audit readiness, and security control assessment.",
            "Security investment governance through risk reduction, ROI/TCO, RTO/RPO, SLA, business continuity, and measurable value realization.",
            "Board reporting, audit engagement, customer security reviews, crisis governance, and executive stakeholder management."
        ]),
        ("Digital Trust, Data Governance & Auditability", [
            "Security governance for critical business systems, confidential records, sensitive data, intellectual property, enterprise workflows, and digital platforms.",
            "System inventory, business-criticality classification, risk assessment, security control baselines, evidence management, and security operating procedures.",
            "Least privilege, segregation of duties, access recertification, logging, audit-trail protection, encryption, backup, recovery, and change-control governance.",
            "Data lifecycle controls covering classification, access, retention, archival, deletion, integrity, traceability, availability, and recoverability.",
            "Security architecture and supplier assurance for ERP, CRM, HRIS, finance, document management, workflow, data platforms, collaboration tools, and SaaS ecosystems."
        ]),
        ("Identity, Zero Trust & Data Protection", [
            "Enterprise IAM, SSO, MFA, conditional access, privileged-access governance, PAM concepts, least privilege, joiner-mover-leaver lifecycle controls, and third-party identity management.",
            "Data classification, DLP, encryption, key management, endpoint protection, secure collaboration, CASB/SSE concepts, and cross-border data governance.",
            "Protection of IP, confidential R&D information, customer and partner data, financial records, commercial information, and internal corporate assets."
        ]),
        ("Cloud, SaaS & Enterprise Application Security", [
            "AWS, Azure, and GCP security governance, landing zones, centralized logging, configuration assurance, CSPM/CNAPP-aligned controls, segmentation, WAF/WAAP, API security, backup, and disaster recovery.",
            "Secure governance of ERP/SAP, CRM, HRIS, finance applications, data platforms, collaboration tools, workflow systems, and business-critical SaaS services.",
            "Cloud and SaaS vendor due diligence, shared-responsibility assessments, architecture review, contractual security controls, continuous assurance, and exit/continuity planning.",
            "Multi-region and multi-zone architecture design for critical services and cross-region disaster recovery."
        ]),
        ("Security Operations, Resilience & Incident Response", [
            "SOC/MDR operating model, SIEM, EDR/XDR, centralized logging, detection engineering, threat intelligence, vulnerability management, incident response, and executive incident reporting.",
            "Ransomware readiness, immutable backup, disaster recovery, business continuity, crisis communication, tabletop exercises, adversary simulation, and recovery governance.",
            "Security observability using ELK Stack, Graylog, Prometheus, LibreNMS, Spiceworks, cloud-native monitoring, and application/system telemetry.",
            "Security metrics covering MTTR, detection coverage, incident severity, response effectiveness, recovery capability, and control effectiveness."
        ]),
        ("DevSecOps, Application & Supply-Chain Security", [
            "Secure SDLC, threat modeling, SAST, DAST, SCA, secrets management, API security, secure CI/CD, policy-as-code, Terraform/IaC governance, Kubernetes/GKE security, and SBOM.",
            "Software-supply-chain security, open-source component governance, secure code review, vulnerability remediation, and application security requirements.",
            "Security governance for internal engineering teams, outsourced development partners, system integrations, APIs, data pipelines, and third-party platforms."
        ]),
        ("AI Governance & Secure Innovation", [
            "ISO/IEC 42001-based AI governance, Responsible AI policies, approved use-case governance, model and vendor risk assessments, data classification, and AI lifecycle management.",
            "GenAI data-leakage prevention, shadow-AI control, secure RAG knowledge-base governance, access control, prompt/data protection, monitoring, auditability, and human oversight.",
            "AI-assisted risk scoring, anomaly detection, threat correlation, and security analytics."
        ]),
        ("Leadership & Stakeholder Management", [
            "Cross-functional and cross-border leadership across Taiwan, China, Asia-Pacific, Europe, and global managed-service partners.",
            "Leadership of cybersecurity, IT operations, cloud operations, infrastructure, SRE, DevOps, applications, service desk, and external service-provider teams.",
            "Multi-million-dollar annual IT and cybersecurity budget responsibility.",
            "Chinese and English executive communication; board reporting; audit facilitation; customer due diligence; vendor governance; and incident crisis leadership."
        ])
    ]
    for d_title, bullets in comp_data:
        add_subheading(doc, d_title)
        for b in bullets:
            add_bullet(doc, b, bold_prefix_colon=False)

    add_heading_1(doc, "Professional Experience")
    
    # 1. Confidential Group
    add_exp_item(
        doc,
        role="Vice President / Group Cybersecurity & Digital Transformation Lead",
        group_note="(Kaohsiung, Shanghai, Taipei / Shengxin, Shengji Network / China Poker City / Taipei Finger International Co., Ltd.) Listed Affiliate Group",
        company_line="Confidential Group | Taiwan | Multi-site Operations",
        period="May 2025 – Present",
        desc="Acting in a CISO-level capacity, leading group cybersecurity strategy, cloud and infrastructure governance, enterprise resilience, and AI governance across multiple business units.",
        scopes=[
            "Group-wide IT and cybersecurity governance across multiple business units and operating sites.",
            "Leadership of approximately 35–40 professionals across cybersecurity, infrastructure, applications, operations, and external service partners.",
            "Annual IT and cybersecurity budget responsibility of approximately USD 10–12 million."
        ],
        achs=[
            "Established group-wide cybersecurity governance based on ISO 27001 and NIST CSF, elevating risk, investment prioritization, resilience metrics, and accountability to board-level oversight.",
            "Built an enterprise cybersecurity risk-management system covering asset criticality, threats, vulnerabilities, control maturity, third-party exposure, and quantified risk matrices.",
            "Developed multi-year cybersecurity roadmaps linked to annual budgets, OGSM objectives, business priorities, and measurable risk-reduction outcomes.",
            "Designed Zero Trust-aligned identity, endpoint, network, cloud, and access-control architecture for critical enterprise applications, data platforms, and partner collaboration.",
            "Implemented centralized SOC/SIEM, EDR, log aggregation, security monitoring, and incident-response procedures, reducing major security incidents by approximately 30% and improving detection and response performance.",
            "Established multi-cloud security observability and configuration-assurance dashboards, improving executive visibility, audit evidence, customer due-diligence response, and remediation governance.",
            "Introduced ISO/IEC 42001-based AI governance, including AI-use policies, risk assessment, data-protection requirements, model lifecycle controls, and third-party AI-service governance.",
            "Secured executive sponsorship and investment for Zero Trust, SOC/SIEM, disaster recovery, and cloud-security initiatives through ROI/TCO, RTO/RPO, SLA, risk-reduction, and audit-readiness analysis.",
            "Coordinated security and digital-transformation initiatives among IT, finance, legal, HR, engineering, business operations, and external service partners."
        ],
        leaving="Seeking a northern Taiwan-based executive opportunity aligned with long-term family, career, and group-level cybersecurity leadership objectives."
    )

    # 2. GameSparcs
    add_exp_item(
        doc,
        role="IT Director",
        group_note="Longzhong Network Co., Ltd. / GameSparcs (Headquartered in Taichung with offices in Taipei, Los Angeles, Sydney, Malta, Hangzhou, Chengdu / XSGames, Longzhong, Wanguo, Haiyu, VIVIDGAMING, Galaxy, Haotian, Jingqi) Publicly Listed Company",
        company_line="Publicly Listed Global Gaming Platform Operator | Taichung, Taiwan",
        period="September 2022 – April 2025",
        desc="Led global cloud architecture, cybersecurity governance, digital-platform resilience, DevOps productivity, and high-availability operations for a publicly listed online platform serving millions of users.",
        scopes=[
            "Led approximately 30–35 professionals across cloud operations, DevOps, SRE, cybersecurity, infrastructure, and application-support functions.",
            "Accountable for multi-cloud infrastructure and cybersecurity budgets of approximately USD 12–14 million annually.",
            "Supported high-traffic, customer-facing digital platforms requiring 24/7 availability, large-scale transaction processing, and resilient cross-region operations."
        ],
        achs=[
            "Designed and delivered a GKE-centered, multi-zone, multi-region architecture with cross-region disaster recovery, enabling 24/7 service availability and a zero-downtime record for critical services.",
            "Built centralized observability, APM, SIEM, cloud logging, monitoring, and incident-response processes, reducing MTTR by approximately 30% through real-time event correlation and improved operational visibility.",
            "Implemented ISO 27001- and NIST CSF-aligned cybersecurity policies and controls for IAM, asset management, secure development, cloud operations, incident response, vendor governance, and audit evidence.",
            "Introduced Zero Trust principles through IAM, SSO, MFA, least privilege, and controlled remote access, strengthening protection for distributed teams, privileged users, and external partners.",
            "Strengthened the external attack surface through API security, WAF/WAAP, DDoS protection, anti-bot controls, rate limiting, API authorization, and network-security governance.",
            "Established incident-response playbooks, security exercises, adversary simulations, crisis communication procedures, and recovery governance, improving organizational readiness and reducing breach impact.",
            "Embedded SAST, DAST, secrets management, CI/CD controls, IaC governance, container-security practices, and software-supply-chain controls into engineering and deployment workflows.",
            "Built FinOps practices across multi-cloud environments, achieving approximately 30% cloud-cost optimization while maintaining performance, resilience, and customer experience during traffic spikes.",
            "Received a Cloud Architecture Excellence Award for a GKE-based multi-cloud platform combining high availability, operational resilience, and cost optimization."
        ],
        leaving="Seeking a senior cybersecurity leadership role with broader group-level governance responsibility and improved alignment with long-term career and lifestyle objectives."
    )

    # 3. Hongen Technology
    add_exp_item(
        doc,
        role="IT Director",
        group_note="Hongen Technology (Headquartered in Shenzhen with offices in Hsinchu, Hangzhou, Chengdu / 泓晏科技)",
        company_line="Electronics Manufacturing & Technology Services | Taiwan and China",
        period="March 2018 – August 2022",
        desc="Led enterprise IT, cybersecurity, data governance, digital transformation, and multi-site technology operations across Taiwan and China.",
        scopes=[
            "Managed approximately 30–40 professionals across IT, cybersecurity, applications, infrastructure, operations, and technology services.",
            "Held annual IT and cybersecurity budget responsibility of approximately USD 5–7 million.",
            "Supported multi-site business operations, R&D, supply-chain, engineering, finance, procurement, and external-partner environments."
        ],
        achs=[
            "Established ISO 27001-aligned cybersecurity governance, risk assessment, asset management, identity and access controls, segmentation, vulnerability management, and auditability for enterprise and operational environments.",
            "Designed secure architecture for ERP, PLM, supply-chain, workflow, analytics, and enterprise systems, improving availability, traceability, data protection, and continuity.",
            "Led secure integration of ERP, MES, PLM, APS, HR, and analytics platforms, enabling controlled information flow, authentication, encryption, audit logging, and management visibility across business functions.",
            "Built data-governance and BI capabilities supporting inventory management, anomaly analysis, operational decision-making, and post-merger data integration.",
            "Developed risk-based vulnerability, patch-management, incident-response, backup, and continuity processes that balanced cybersecurity improvement with operational requirements.",
            "Coordinated cybersecurity and resilience initiatives across IT, engineering, operations, procurement, finance, and external service partners in Taiwan and China.",
            "Supported organizational integration and data-asset due diligence during acquisition by a major Chinese manufacturing group."
        ],
        add_ctx="Prior IT/OT convergence experience provides practical capability for complex infrastructure, supply-chain, partner, and operational integration needs.",
        leaving="Following acquisition by a major Chinese conglomerate, pursued opportunities aligned with broader technology and cybersecurity leadership goals."
    )

    # 4. HyWeb
    add_exp_item(
        doc,
        role="IT Manager",
        group_note="HyWeb Technology Co., Ltd. (Headquartered in Hsinchu with offices in Taipei, Taichung, Kaohsiung, Thailand, Beijing / 凌網科技、凌網知識) Publicly Listed Company",
        company_line="Publicly Listed IT Services Company | Taiwan, China and Thailand",
        period="May 2014 – February 2018",
        desc="Led regional IT operations, service management, data-center strategy, infrastructure governance, and cross-border technology support for offices in Taipei, Hsinchu, Taichung, Beijing, and Thailand.",
        scopes=[
            "Managed approximately 34–45 professionals across infrastructure, enterprise applications, service desk, support, and operations.",
            "Led regional IT budgets, data-center operations, backup architecture, disaster-recovery planning, and service-management improvement."
        ],
        achs=[
            "Implemented ITSM and cross-regional service-management processes, improving service consistency, accountability, incident response, and operational transparency across Asia-Pacific offices.",
            "Established regional data-center strategy, high-availability architecture, backup governance, and disaster-recovery plans to support critical business services.",
            "Introduced PMP and agile delivery practices, improving cross-border project governance, delivery predictability, and stakeholder collaboration.",
            "Strengthened infrastructure standards, security operations, and technology-service governance for geographically distributed operations."
        ]
    )

    # 5. Kuang Nan Group
    add_exp_item(
        doc,
        role="IT Manager / Digital Transformation Lead",
        group_note="Kuang Nan Group / Fengchen Group – Aiju Computer (Headquartered in Taichung with offices in Taipei, Shenzhen, Shanghai, Zhangjiagang / 光南集團) Publicly Listed Company",
        company_line="Retail, Distribution & Technology Services Group | Taiwan and China",
        period="July 2011 – April 2014",
        desc="Led enterprise-system integration, business-data governance, and digital transformation during merger-and-acquisition activities for a retail and distribution group.",
        scopes=[
            "Managed approximately 15–20 IT professionals.",
            "Responsible for ERP, POS, CRM, enterprise data integration, business continuity, application governance, and transformation budgets."
        ],
        achs=[
            "Consolidated ERP, POS, CRM, and analytics platforms to improve group-level visibility into sales, inventory, customer behavior, and operational performance.",
            "Led data cleansing, master-data governance, migration, and business-system consolidation to maintain continuity during cross-border M&A.",
            "Enabled management reporting and decision support through integrated enterprise-data architecture."
        ]
    )

    # 6. Borland
    add_exp_item(
        doc,
        role="IT Manager",
        group_note="Borland (美商 寶藍) Multinational Company",
        company_line="Multinational Software Company | Taiwan and China",
        period="March 2008 – June 2011",
        desc="Built cross-border IT collaboration platforms, high-availability enterprise systems, infrastructure services, and data-protection capabilities for Taiwan and China R&D and operations teams.",
        scopes=None,
        achs=[
            "Designed high-availability database, application, backup, and recovery architecture for enterprise applications and mission-critical transaction environments.",
            "Supported deployment of software-lifecycle-management platforms, engineering collaboration systems, and regional IT services.",
            "Maintained operational stability, data protection, and cross-border collaboration capability in multinational operating environments."
        ]
    )

    # 7. Sybase
    add_exp_item(
        doc,
        role="IT Manager",
        group_note="Sybase (美商 賽貝斯) Multinational Company",
        company_line="Multinational Database and Enterprise Software Vendor | Asia-Pacific",
        period="August 2002 – February 2008",
        desc="Led Asia-Pacific data-center consolidation, infrastructure standardization, database availability, regional IT operations, and mission-critical support for telecommunications and financial-services customers.",
        scopes=[
            "Managed regional infrastructure teams of approximately 10–15 professionals.",
            "Supported enterprise database, high-availability, failover, backup, and data-center environments."
        ],
        achs=[
            "Implemented unified Active Directory, data-center consolidation, high-availability database, and failover strategies across Asia-Pacific operations.",
            "Delivered carrier-grade database tuning, availability, recovery, and performance solutions for business-critical customer environments.",
            "Ensured operational continuity and audit compliance during regional consolidation and global customer projects."
        ]
    )

    add_heading_1(doc, "Professional Certifications")
    for cert in [
        "ISO 27001 Information Security Management System — Lead Auditor / Internal Auditor",
        "ISO/IEC 42001 AI Management System — AI Governance",
        "Project Management Professional — PMP",
        "Certified ScrumMaster — CSM",
        "Oracle Certified Professional — OCP",
        "Oracle Certified Associate — OCA",
        "ESG Sustainability Planner"
    ]:
        add_bullet(doc, cert, bold_prefix_colon=False)

    add_heading_1(doc, "Academic Education")
    edu_list = [
        ("Ph.D. in Information Technology Management", "Chaoyang University of Technology, College of Informatics\nSeptember 2009 – June 2013", "Research interests: Cybersecurity governance, cloud security, AI governance, data governance, software engineering, DevSecOps, and digital transformation strategy."),
        ("Master of Science in Information Science", "National Chung Hsing University, Institute of Computer Science\nSeptember 2004 – June 2006", "Focus areas: Modern cryptography, data mining, data governance, software engineering, systems analysis, and enterprise information systems."),
        ("Bachelor of Computer Science", "National Chung Hsing University, Department of Computer Science\nSeptember 2000 – June 2004", "Focus areas: Software development lifecycle, software engineering, systems analysis, enterprise architecture, and application development.")
    ]
    for d, s, desc in edu_list:
        add_subheading(doc, d, size=Pt(10.5))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(s)
        set_run_font(r, size=Pt(9.5), bold=True, color=COLOR_SECONDARY)
        add_body_p(doc, desc, space_after=4)

    add_heading_1(doc, "Awards, Keynotes, Publications & Speaking")
    
    add_subheading(doc, "1. Keynote Addresses & Professional Speaking", size=Pt(11))
    add_media_entry(doc, "MongoDB.local Taipei Keynote Speaker", "MongoDB Inc. | September 2024", "Delivered Keynote on global multi-region high-availability data architecture, sharing how MongoDB Atlas fully managed cloud database supported global cross-region low-latency gaming with 99.995% availability.", url="https://www.ithome.com.tw/pr/163534", image_filename="2024_MongoDB02.jpg", img_width=Inches(3.4))
    add_media_entry(doc, "CIO Taiwan Value Academy 17th Keynote Speaker", "CIO Taiwan Value Academy | May 2024", "Keynote presentation to enterprise CIOs/CISOs on multi-cloud resource governance, cybersecurity compliance, FinOps TCO optimization, and architecture resilience.", url="https://www.cio.com.tw/events/value-index-2/", image_filename="問題比答案重要.png", img_width=Inches(3.2))
    add_media_entry(doc, "Executive Yuan DGPA Master Lecture Series", "Directorate-General of Personnel Administration | June 2024", "Invited master instructor for civil servants, promoting cloud modernization, Zero Trust architecture, and resilient IT governance.", url="https://elearn.hrd.gov.tw/info/10042804", image_filename="e等公務園__洞見機遇-雲端管理與實務01.png", img_width=Inches(3.0))
    add_media_entry(doc, "Google Cloud Official Global Customer Success Story", "Google Cloud APAC Video | March 2024", "Featured in official video case study on deploying Kubernetes/GKE multi-zone architecture achieving 100% Zero Outage during peak traffic spikes.", url="https://youtu.be/_kTZSZ_0lNE?si=CT2lo8c4IF0zI1Ki")

    add_subheading(doc, "2. Academic Publications & Research", size=Pt(11))
    add_media_entry(doc, "Springer SCI Journal Paper Publication: Reversible secret-image sharing with high visual quality", "Multimedia Tools and Applications (Springer Nature, Vol. 74) | June 2014", "Authors: Ching-Chiuan Lin, Lun-Hao Liao (Howard Liao), Kuo-Feng Hwang, Shih-Chieh Chen. Proposes a high visual quality reversible secret-image sharing scheme in cryptographic image security.", url="https://link.springer.com/article/10.1007/s11042-012-1190-1")
    add_media_entry(doc, "International Journal Publication: Utilizing GIS and GPS in Designing a Trilingual Tourist APP", "Applied Science and Management Research (Vol. 2) | May 2015", "Research on location-based mobile systems integrating GIS and GPS technologies for multilingual travel platforms.")
    add_media_entry(doc, "National Central Library Legal Archive & Academic Advising", "National Central Library (Advising Archive ID: 106IKTC0183002) | October 2014", "Ph.D. dissertation on ITIL V3 and SOA framework governance. Registered research advisor in National Central Library under ID 106IKTC0183002.", url="https://ndltd.ncl.edu.tw/")

    add_subheading(doc, "3. Media Coverage & Technical Features", size=Pt(11))
    add_media_entry(doc, "CIO Taiwan Magazine Cover Feature Interview", "CIO Taiwan Magazine (Issue 2024.05) | May 2024", "In-depth interview by Editor-in-Chief on Multi-Cloud Networking (MCN), Kubernetes native architecture, and FinOps cost optimization.", url="https://www.cio.com.tw/interview-howard-liao-online-gamesparcs-it-director/", image_filename="2024_CIO報導.png", img_width=Inches(3.0))
    add_media_entry(doc, "iThome Tech Feature Report", "iThome Computer Daily | September 2024", "Full-page technology feature highlighting Howard Liao's data platform architecture delivering 99.995% availability for global operations.", url="https://www.ithome.com.tw/pr/163534", image_filename="IThome_MangoDB.png", img_width=Inches(3.2))
    add_media_entry(doc, "DIGITIMES Tech News Feature: MongoDB 8.0 Launch", "DIGITIMES Tech News | September 2024", "News coverage citing GameSparcs architecture cases in generative AI and real-time database workloads.", url="https://www.digitimes.com.tw/tech/dt/n/shwnws.asp?id=0000704377_KI14VVDK6CQ8TB8OXC3GJ")

    add_subheading(doc, "4. Awards & Community Leadership", size=Pt(11))
    add_bullet(doc, "Cloud Architecture Excellence Award — Recognized for a GKE-based multi-cloud architecture delivering high availability, operational resilience, and approximately 30% cloud-cost optimization.", bold_prefix_colon=False)
    add_bullet(doc, "Keynote speaker at cybersecurity summits, sharing practical experience in hybrid-cloud security, Zero Trust implementation, ISO 27001, ISO/IEC 42001, cybersecurity governance, and digital resilience.", bold_prefix_colon=False)
    add_bullet(doc, "Invited speaker at CIO academies and technology communities on multi-cloud governance, AI governance, enterprise digital transformation, and cybersecurity strategy.", bold_prefix_colon=False)
    add_media_entry(doc, "Cancer NoNo Foundation Board Member & CIO", "Cancer NoNo Foundation (CSR) | 2022 – Present", "Serving as Board Director and CIO, leading CSR digital initiatives, children dietary education, and secure medical IT platforms.", url="https://www.cancer-nono.org.tw/web/about/page.php?lang=zh_tw&scid=7&sid=5")

    out_file = os.path.join(target_dir, "Howard_Liao_CISO_Resume_EN.docx")
    doc.save(out_file)
    print(f"Generated: {out_file}")

# 2. GENERATE TRADITIONAL CHINESE (Howard_Liao_CISO_Resume_ZH.docx)
def build_zh_docx():
    doc = create_base_doc()
    add_header(
        doc,
        name_text="廖倫豪 博士 (Howard Liao, Ph.D.)",
        title_text="集團資安長 暨 科技副總 (Group CISO)",
        subtitle_text="全球資安治理、數位信任與架構韌性 (對 董事會、董事長、總經理、CEO 專用履歷)",
        contact_text1="台灣 (支援全球跨國據點)  |  行動電話：+886-975-323161  |  電子郵件：Liao.Howard@gmail.com",
        contact_text2="LinkedIn 領英：linkedin.com/in/howardliao78  |  作品集網站：https://howardliao.github.io/portfolio/",
        east_asia="Microsoft JhengHei"
    )
    
    add_heading_1(doc, "高階主管職涯定位 (Executive Profile)", east_asia="Microsoft JhengHei")
    add_body_p(doc, "CISO 級別高階科技主管，具備 27+ 年企業 IT 領導經驗、15+ 年資訊安全實務，以及 10+ 年在上市櫃公司、跨國集團與受監管業務環境中主導資安戰略、雲端治理、數位韌性與企業數位轉型的深厚底蘊。")
    add_body_p(doc, "兼具董事會層級之資安治理高度，以及零信任 (Zero Trust)、身分存取管理 (IAM)、多雲資安 (Multi-Cloud Security)、SOC/SIEM/EDR、資安事件應變 (Incident Response)、DevSecOps、資料保護、第三方風險管理與 AI 治理 (ISO 42001) 的實戰技術深度。")
    add_body_p(doc, "精通保護關鍵業務系統、敏感機密紀錄、核心智慧財產權、企業級應用程式、雲端/SaaS 平台與跨境資料流。以務實、風險驅動的方法持續提升資安成熟度、營運韌性、稽核整備度、災難復原力與數位信任。")
    add_body_p(doc, "擁有卓越的成果紀錄，能將資安風險與技術優先級轉化為董事會與高階主管重視的商業價值、財務效益、營運指標與法規遵循結果。建構具備高擴展性的資安營運模型，在保護關鍵資產與強化業務連續性的同時，全力支援集團之跨國擴張。")
    add_body_p(doc, "過去在 IT/OT 融合與跨國多廠區工業製造環境之豐富實務，進一步奠定了駕馭複雜基礎架構、營運韌性、供應鏈整合與合作夥伴生態系安全防護的全面能力。")

    add_heading_1(doc, "董事會與高階領導價值主張 (Leadership Value Proposition)", east_asia="Microsoft JhengHei")
    add_bullet(doc, "CISO 董事會戰略治理：透過風險矩陣、RTO/RPO、SLA、ROI/TCO 與成熟度藍圖，將資安態勢、風險暴露、事件與控制落差精準轉化為具體商業衝擊與決策依據。")
    add_bullet(doc, "集團級資安營運架構：跨越企業總部、國際據點、多公有雲環境、SaaS 平台、託管服務商 (MSP) 與商業夥伴，建立全集團標準化資安營運體系。")
    add_bullet(doc, "數位信任與架構韌性：透過身分治理、安全監控、加密、完整稽核軌跡、備份、復原與即時應變機制，全方位守護關鍵系統、機密紀錄與智慧財產權。")
    add_bullet(doc, "雲原生技術深度：結合高階策略與 AWS、Azure、GCP、Kubernetes/GKE、Zero Trust、API 安全、WAF/WAAP、SIEM/SOC、EDR/XDR 與 IaC 自動化之深厚實作經驗。")
    add_bullet(doc, "AI 治理 (ISO/IEC 42001)：落地 ISO/IEC 42001 AI 管理系統標準，在推動企業負責任採用 GenAI 的同時，嚴密控管 Shadow AI、資料外洩、IP 侵權與第三方模型風險。")
    add_bullet(doc, "業務賦能與成長思維：與業務、財務、法務、人資、研發、營運、產品與 IT 團隊緊密協同，將資訊安全無縫嵌入業務增長、數位轉型與客戶信任之中。")

    add_heading_1(doc, "八大核心職能與技術治理體系 (Core Competencies)", east_asia="Microsoft JhengHei")
    comp_data_zh = [
        ("資安策略、治理與風險管理 (Cybersecurity Strategy, Governance & Risk)", [
            "集團資安策略、營運模型、政策架構、風險胃納、成熟度評估、KPI/KRI 指標設計、高階主管報告與多年期資安藍圖規劃。",
            "ISO 27001/27002、NIST CSF、Zero Trust 零信任架構、風險管理、供應商風險、資安盡職調查、稽核整備度與控制措施評估。",
            "藉由風險降低、ROI/TCO 分析、RTO/RPO、SLA、業務連續性與可量化價值實現，落實資安投資治理。",
            "董事會定期匯報、外部稽核應對、重要客戶資安審查、危機治理與高階利害關係人管理。"
        ]),
        ("數位信任、數據治理與可稽核性 (Digital Trust, Data Governance & Auditability)", [
            "關鍵業務系統、機密檔案、敏感資料、智慧財產權、企業工作流與數位平台的全面資安治理。",
            "資產盤點、業務關鍵性分級、風險評估、安全控制基準線、稽核佐證管理與標準作業程序 (SOP)。",
            "最小權限、職責分離 (SoD)、存取定期複核、日誌記錄、稽核軌跡保護、資料加密、備份復原與變更控制治理。",
            "涵蓋資料分類、存取、留存、封存、銷毀、完整性、可追溯性、可用性與可復原性之數據生命週期管理。",
            "針對 ERP、CRM、HRIS、財務系統、文件管理、資料中台、協作工具與 SaaS 生態系的安全架構與供應商審查。"
        ]),
        ("身分鑑別、零信任與資料保護 (Identity, Zero Trust & Data Protection)", [
            "企業級 IAM、SSO、MFA、條件式存取、特權存取治理 (PAM)、最小權限、員工進轉離生命週期控管與第三方身分管理。",
            "資料分級分類、DLP 防外洩、加密與金鑰管理、端點安全、安全協作、CASB/SSE 概念與跨境資料流動治理。",
            "全面防護核心智慧財產權、機密研發資料、客戶與合作夥伴資訊、財務紀錄、商業情報與內部核心資產。"
        ]),
        ("多雲、SaaS 與企業應用安全 (Cloud, SaaS & Enterprise Application Security)", [
            "AWS、Azure 與 GCP 資安治理、Landing Zones、集中化日誌、組態合規、CSPM/CNAPP 控制項、網路微隔離、WAF/WAAP、API 安全、備份與跨區災難復原。",
            "ERP/SAP、CRM、HRIS、財務軟體、資料平台、協作工具、工作流與關鍵 SaaS 服務之安全管控。",
            "雲端與 SaaS 供應商盡職調查、共同責任模型評估、架構審查、合約安全條款、持續合規保證與退場/業務連續性計畫。",
            "針對關鍵服務之跨區域 (Multi-Region) 與跨可用區 (Multi-Zone) 高可用架構設計與跨區災備機制。"
        ]),
        ("維運監控、營運韌性與事件應變 (Security Operations, Resilience & Incident Response)", [
            "SOC/MDR 營運模型、SIEM、EDR/XDR、集中式日誌聚合、偵測工程、威脅情資、弱點管理、事件應變與高階事件通報機制。",
            "勒索軟體防禦、不可變備份 (Immutable Backup)、災難復原、業務連續性、危機通訊、兵棋推演、紅藍對抗演練與復原治理。",
            "運用 ELK Stack、Graylog、Prometheus、LibreNMS、Spiceworks、雲端原生監控與系統遙測落實全域可觀測性。",
            "追蹤 MTTR、偵測覆蓋率、事件嚴重度、應變處置成效、復原能力與控制措施有效性等量化指標。"
        ]),
        ("DevSecOps、應用程式與軟體供應鏈安全 (DevSecOps, Application & Supply-Chain Security)", [
            "安全軟體開發生命週期 (SSDLC)、威脅建模、SAST、DAST、SCA、機敏密鑰管理、API 安全、安全 CI/CD、Policy-as-Code、IaC 治理、Kubernetes/GKE 資安與 SBOM。",
            "軟體供應鏈安全、開源元件合規治理、安全代碼審查、漏洞修補流程與應用程式安全規範要求。",
            "針對內部研發團隊、委外開發廠商、系統整合商、API 串接、數據管道與第三方平台之全面安全管理。"
        ]),
        ("AI 治理與安全創新 (AI Governance & Secure Innovation)", [
            "基於 ISO/IEC 42001 之 AI 治理體系、負責任 AI 政策、核准用例管理、模型與供應商風險評估、資料分級與 AI 生命週期控管。",
            "生成式 AI 資料防洩漏 (GenAI DLP)、Shadow AI 管控、安全 RAG 知識庫治理、存取控制、Prompt/資料保護、即時監控、可稽核性與人機複核 (Human-in-the-Loop)。",
            "結合 AI 輔助風險評分、異常行為偵測、威脅事件關聯分析與進階資安數據分析。"
        ]),
        ("高階領導力與跨域利害關係人管理 (Leadership & Stakeholder Management)", [
            "具備橫跨台灣、中國大陸、亞太地區、歐洲與全球託管服務夥伴的跨國跨職能領導實務。",
            "統領資安、IT 維運、雲端維運、基礎架構、SRE、DevOps、應用程式開發、服務台與外部廠商團隊。",
            "管控每年高達數百萬至千萬美元之全球 IT 與資安預算。",
            "中英文高階溝通能力；董事會定期報告；各類稽核應對；重要客戶調查評估；供應商管理；資安危機處置領導力。"
        ])
    ]
    for d_title, bullets in comp_data_zh:
        add_subheading(doc, d_title)
        for b in bullets:
            add_bullet(doc, b, bold_prefix_colon=False)

    add_heading_1(doc, "專業歷練與職涯成就里程碑 (Professional Experience)", east_asia="Microsoft JhengHei")
    zh_labels = {"scope": "領導範疇 (Leadership Scope)", "ach": "重大成就與量化效益 (Selected Achievements & Impact)", "ctx": "補充背景 (Additional Context)", "leaving": "離職原因 (Reason for Leaving)"}
    
    # 1. 關聯集團
    add_exp_item(
        doc,
        role="集團副總 暨 資安與數位轉型負責人 (Vice President)",
        group_note="(高雄、上海、台北 / 盛欣、盛碁網絡/中國 波克/台北 芬格國際有限公司) 上市櫃公司 關聯集團",
        company_line="跨國集團 | 台灣 | 多據點運營",
        period="2025.05 – 至今",
        desc="以 CISO 級別角色主導全集團資訊安全戰略、多雲與基礎架構治理、企業營運韌性及 AI 治理，橫跨多個事業群與跨國營運據點。",
        scopes=[
            "全集團跨事業部、跨營運據點之 IT 與資安最高治理架構。",
            "帶領約 35–40 位專業人員，涵蓋資訊安全、基礎架構、應用系統、維運與外部託管服務夥伴。",
            "管控每年約 1,000–1,200 萬美元之全集團 IT 與資安總預算。"
        ],
        achs=[
            "建立基於 ISO 27001 與 NIST CSF 之集團資安治理框架，將風險管理、投資排序、韌性指標與權責機制提升至董事會層級監管。",
            "打造企業級資安風險管理體系，涵蓋資產關鍵性分級、威脅情資、弱點管理、控制措施成熟度、第三方風險與量化風險矩陣。",
            "制定緊密連結年度預算、OGSM 目標、商業優先級與可量化風險降減成果之多年期資安藍圖。",
            "針對關鍵企業應用、資料中台與夥伴協作，設計符合零信任 (Zero Trust) 原則之身分、端點、網路、雲端與存取控制架構。",
            "導入集中式 SOC/SIEM、EDR、日誌聚合、資安即時監控與應變標準作業流程，降低重大資安事件約 30%，大幅提升偵測處置效能。",
            "建立多雲資安可觀測性與組態合規儀表板，提升高階主管能見度、稽核佐證效率、重要客戶盡職調查回應與弱點修復治理。",
            "推動導入 ISO/IEC 42001 AI 管理系統，訂定 AI 使用規範、風險評估、資料防護標準、模型生命週期控制與第三方 AI 服務治理。",
            "透過 ROI/TCO、RTO/RPO、SLA、風險降減幅度與稽核整備度分析，成功爭取零信任、SOC/SIEM、災難復原與雲端資安等重大投資贊助。",
            "有效協調 IT、財務、法務、人資、研發、業務營運與外部服務供應商，推進資安與數位轉型專案。"
        ],
        leaving="尋求位於台灣北部、能與長期家庭生活規劃、職涯發展及集團級資安長治理目標高度契合的高階主管機會。",
        labels=zh_labels
    )

    # 2. 隆中網絡
    add_exp_item(
        doc,
        role="資訊處長 / IT Director",
        group_note="隆中網絡股份有限公司 / GameSparcs (總部在台中，並在台北、洛杉磯、雪梨、馬爾他、杭州與成都設有據點 / 向上國際XSGames、隆中網絡、萬國遊戲、海淯遊戲、VIVIDGAMING、銀河網絡、浩天遊戲、晶綺科技) 上市櫃公司",
        company_line="上市櫃全球線上娛樂與遊戲平台營運商 | 台中, 台灣",
        period="2022.09 – 2025.04",
        desc="主導全球雲端架構、資安治理、數位平台韌性、DevOps 效能與高可用性維運，支撐服務全球數百萬玩家之上市櫃線上平台。",
        scopes=[
            "帶領約 30–35 位專業技術人員，涵蓋雲端維運、DevOps、SRE、資訊安全、基礎架構與應用系統支援。",
            "負責每年約 1,200–1,400 萬美元之多雲基礎架構與資安預算。",
            "全面支撐需 24/7 全天候高可用、巨量交易處理與跨區域高韌性之高並發大型面向客戶數位平台。"
        ],
        achs=[
            "設計並交付以 GKE 為核心之多可用區、多區域架構與跨區災備機制，確保 24/7 服務不中斷，創下關鍵服務 100% 零停機 (Zero Outage) 紀錄。",
            "建立集中式可觀測性、APM、SIEM、雲端日誌聚合、即時監控與應變機制，透過即時事件關聯分析將 MTTR 縮短約 30%。",
            "落實符合 ISO 27001 與 NIST CSF 之資安政策與控制措施，涵蓋 IAM、資產管理、安全開發、雲端維運、事件應變、供應商治理與稽核軌跡。",
            "透過 IAM、SSO、MFA、最小權限與安全遠端存取導入零信任原則，強化分散式團隊、特權用戶與外部合作夥伴之防護力。",
            "強化外部攻擊面防禦，導入 API 安全防護、WAF/WAAP、DDoS 緩解、防爬蟲機制、速率限制、API 鑑權與網路安全治理。",
            "建立資安應變 Playbooks、應變演練、紅藍對抗推演、危機通訊程序與復原治理，大幅提升組織資安整備度並降低資安事件衝擊。",
            "將 SAST、DAST、機敏密鑰管理、CI/CD 安全控制、IaC 治理、容器安全與軟體供應鏈安全無縫整合至工程研發與部署流程。",
            "推動多雲 FinOps 雲端財務營運實踐，在業務流量劇增期間維持極致效能與穩定性的同時，達成約 30% 雲端成本優化。",
            "憑藉兼具超高可用性、營運韌性與成本最佳化之 GKE 多雲架構，榮獲「雲端架構卓越獎 (Cloud Architecture Excellence Award)」。"
        ],
        leaving="尋求具備更廣泛全集團治理權責、能與長期職涯目標更緊密契合的資深資安長領導職位。",
        labels=zh_labels
    )

    # 3. 泓晏科技
    add_exp_item(
        doc,
        role="資訊處長 / IT Director",
        group_note="泓晏科技 (總部在深圳，並在新竹、杭州與成都設有據點)",
        company_line="電子製造與高科技技術服務集團 | 台灣 與 中國大陸",
        period="2018.03 – 2022.08",
        desc="統領企業 IT、資訊安全、數據治理、數位轉型及跨越台灣與中國大陸之多廠區多據點科技營運體系。",
        scopes=[
            "管理約 30–40 位專業人員，橫跨 IT、資安、企業應用、基礎架構、維運與技術服務範疇。",
            "管控每年約 500–700 萬美元之 IT 與資訊安全年度預算。",
            "全面支援多廠區營運、研發中心、供應鏈、工程、財務、採購與外部合作夥伴環境。"
        ],
        achs=[
            "為企業與製造營運環境建立符合 ISO 27001 之資安治理、風險評估、資產盤點、身分與存取控制、網路微隔離、弱點管理與可稽核性。",
            "為 ERP、PLM、供應鏈、工作流、數據分析與企業核心系統設計安全架構，提升高可用性、可追溯性、資料防護與業務連續性。",
            "主導 ERP、MES、PLM、APS、HR 與 BI 分析平台之安全整合，實現跨業務流程之受控資料流、身分鑑別、傳輸加密、稽核日誌與管理能見度。",
            "建置數據治理與商業智慧 (BI) 能力，有效支援庫存智慧管理、異常分析、營運決策與併購後數據資產整合。",
            "建立兼顧資安防護與生產營運連續性之風險驅動漏洞修補、修補管理、事件處置、備份與業務連續性程序。",
            "橫跨台灣與大陸協調 IT、工程、製造營運、採購、財務及外部服務供應商之資安與營運韌性專案。",
            "於集團被中國大型製造業集團收購期間，主導組織 IT 整合與數據資產之盡職調查。"
        ],
        add_ctx="深厚的 IT/OT 融合經驗，為複雜基礎架構、供應鏈整合、夥伴生態系與跨國製造運營提供堅實支撐。",
        leaving="隨公司順利併入大型製造集團後，為追求具備更廣闊戰略發展之資訊安全與科技高階領導機會而離任。",
        labels=zh_labels
    )

    # 4. 凌網科技
    add_exp_item(
        doc,
        role="資訊部經理 / IT Manager",
        group_note="凌網科技股份有限公司 (總部在新竹，並在台北、台中、高雄、泰國、北京設有據點 / 凌網科技、凌網知識) 上市櫃公司",
        company_line="上市櫃大型資訊服務集團 | 台灣、中國大陸 與 泰國",
        period="2014.05 – 2018.02",
        desc="負責亞太區 IT 維運、服務管理 (ITSM)、資料中心戰略、基礎架構治理與跨越台北、新竹、台中、北京與泰國之跨境技術支援。",
        scopes=[
            "管理 34–45 位專業工程師，橫跨基礎架構、企業應用、Service Desk、技術支援與機房維運。",
            "主導區域 IT 預算、資料中心營運、備份架構、災難復原計畫與服務管理流程優化。"
        ],
        achs=[
            "導入 ITSM 與跨區域服務管理機制，提升亞太各分公司之服務一致性、當責性、事件處置速度與營運透明度。",
            "制定區域資料中心戰略、高可用架構、備份治理與災難復原方案，強力支撐關鍵業務服務。",
            "引入 PMP 與敏捷交付實務，強化跨國專案治理、交付可預測性與利害關係人協作效能。",
            "為跨國地理分散據點建立一致化之基礎架構標準、資安營運規範與技術服務體系。"
        ],
        labels=zh_labels
    )

    # 5. 光南集團
    add_exp_item(
        doc,
        role="資訊部經理 / 數位轉型負責人",
        group_note="光南集團 / 峰晨集團 – 艾居電腦 (總部在台中，並在台北、深圳、上海、張家港設有據點 / 光南集團) 上市櫃公司",
        company_line="零售通路、供應鏈物流與科技服務集團 | 台灣 與 中國大陸",
        period="2011.07 – 2014.04",
        desc="於零售與流通集團併購期間，主導企業核心系統整合、商業數據治理與數位轉型推動。",
        scopes=[
            "管理 15–20 位專業 IT 技術團隊成員。",
            "負責 ERP、POS、CRM、企業數據整合、業務連續性、應用系統治理與轉型預算。"
        ],
        achs=[
            "整併 ERP、POS、CRM 與數據分析中台，全面提升集團對門市銷售、庫存周轉、會員行為與營運績效之全局掌控力。",
            "主導跨國併購期間之數據清洗、主數據治理 (MDM)、系統遷移與業務系統整合，確保業務不中斷。",
            "建構整合式企業數據架構，全面賦能管理階層之即時報表分析與策略決策支援。"
        ],
        labels=zh_labels
    )

    # 6. Borland
    add_exp_item(
        doc,
        role="資訊部經理 / IT Manager",
        group_note="美商寶藍 (Borland Taiwan) 跨國外商公司",
        company_line="跨國軟體與生命週期管理外商集團 | 台灣 與 中國大陸",
        period="2008.03 – 2011.06",
        desc="為台灣與中國大陸研發及營運團隊建構跨國 IT 協作平台、高可用企業系統、基礎架構服務與數據保護機制。",
        scopes=None,
        achs=[
            "為企業級應用程式與關鍵交易環境設計高可用資料庫、應用中台、備份與災難復原架構。",
            "支援軟體生命週期管理 (ALM) 平台、工程協作系統與區域 IT 服務之全面部署與維運。",
            "維持跨國運營環境中之高度系統穩定性、資料安全保護與跨境研發協作效能。"
        ],
        labels=zh_labels
    )

    # 7. Sybase
    add_exp_item(
        doc,
        role="資訊部經理 / IT Manager",
        group_note="美商賽貝斯 (Sybase) 跨國外商公司",
        company_line="跨國頂級資料庫與企業軟體外商集團 | 亞太區",
        period="2002.08 – 2008.02",
        desc="主導亞太區資料中心整併、基礎架構標準化、高可用資料庫架構、區域 IT 維運，並為電信與金融級客戶提供關鍵任務技術支援。",
        scopes=[
            "管理亞太區基礎架構團隊約 10–15 位資深工程師。",
            "全面支援企業級大型資料庫、高可用叢集、容錯移轉、備份與資料中心營運。"
        ],
        achs=[
            "於亞太區據點建置統一 Active Directory 目錄服務、資料中心整併、高可用資料庫與自動容錯移轉策略。",
            "為電信運營商與金融機構關鍵任務環境提供電信級 (Carrier-Grade) 資料庫調優、高可用性、災難復原與效能解決方案。",
            "在跨國區域整併與全球大客戶專案執行期間，確保 100% 業務營運連續性與各項稽核法規遵循。"
        ],
        labels=zh_labels
    )

    add_heading_1(doc, "專業國際證照資格 (Certifications)", east_asia="Microsoft JhengHei")
    for cert in [
        "ISO 27001 資訊安全管理系統 — 主任稽核員 / 內部稽核員 (Lead Auditor)",
        "ISO/IEC 42001 人工智慧管理系統 (AIMS) — AI 治理主任稽核員認證",
        "國際專案管理師認證 — PMP (Project Management Professional)",
        "敏捷專案大師認證 — CSM (Certified ScrumMaster)",
        "Oracle 官方認證專業專家 — OCP (Oracle Certified Professional)",
        "Oracle 官方認證專員 — OCA (Oracle Certified Associate)",
        "ESG 永續規劃師證照 (ESG Sustainability Planner)"
    ]:
        add_bullet(doc, cert, bold_prefix_colon=False, east_asia="Microsoft JhengHei")

    add_heading_1(doc, "正規學術學位 (Academic Education)", east_asia="Microsoft JhengHei")
    edu_list_zh = [
        ("資訊科技管理 博士 (Ph.D. in IT Management)", "朝陽科技大學 資訊科技管理研究所 (College of Informatics)\n2009.09 – 2013.06", "研究領域：資訊安全治理、雲端安全、AI 治理、數據治理、軟體工程、DevSecOps 與企業數位轉型戰略。"),
        ("資訊科學 碩士 (Master of Science in Information Science)", "國立中興大學 資訊科學與工程研究所\n2004.09 – 2006.06", "專注領域：現代密碼學 (Modern Cryptography)、資料探勘、數據治理、軟體工程、系統分析與企業資訊系統架構。"),
        ("資訊科學與工程 學士 (Bachelor of Computer Science)", "國立中興大學 資訊科學與工程學系\n2000.09 – 2004.06", "專注領域：軟體開發生命週期 (SDLC)、軟體工程、系統分析、企業架構與應用程式開發。")
    ]
    for d, s, desc in edu_list_zh:
        add_subheading(doc, d, size=Pt(10.5), east_asia="Microsoft JhengHei")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(s)
        set_run_font(r, size=Pt(9.5), bold=True, color=COLOR_SECONDARY, east_asia="Microsoft JhengHei")
        add_body_p(doc, desc, space_after=4, east_asia="Microsoft JhengHei")

    add_heading_1(doc, "獲獎榮譽、大會演講、國際期刊論文與媒體報導", east_asia="Microsoft JhengHei")
    
    add_subheading(doc, "1. Keynote 大會主題演講與實務分享", size=Pt(11), east_asia="Microsoft JhengHei")
    add_media_entry(doc, "MongoDB.local Taipei Keynote 大會主講《運用 MongoDB Atlas 建構全球規模遊戲資料管理平台》", "MongoDB 原廠大會 (Taipei) | 2024.09", "擔任 Keynote 主講人，公開發表全球跨區高可用資料架構，闡述如何藉由 MongoDB Atlas 全託管雲端資料庫兼具高擴展與 99.995% 高可用性，支撐全球百萬玩家跨區即時對弈與高併發讀寫需求。", url="https://www.ithome.com.tw/pr/163534", image_filename="2024_MongoDB02.jpg", img_width=Inches(3.4), east_asia="Microsoft JhengHei", url_label="🔗 查證佐證出處 / URL：")
    add_media_entry(doc, "CIO Taiwan 價值學院第十七屆大會 Keynote 演講《洞見機遇-雲端管理與實務》", "CIO Taiwan 價值學院 (Taipei) | 2024.05", "受邀擔任 Keynote 主講人，向與會上市櫃企業 CIO 與 CISO 分享多雲資源治理、安全合規與 FinOps 實務落地心法，探討雲端現代化治理與架構韌性。", url="https://www.cio.com.tw/events/value-index-2/", image_filename="問題比答案重要.png", img_width=Inches(3.2), east_asia="Microsoft JhengHei", url_label="🔗 查證佐證出處 / URL：")
    add_media_entry(doc, "行政院人事行政總處 e等公務園+學習平臺《洞見機遇-雲端管理與實務》大師講座數位課程", "行政院人事行政總處 e等公務園數位學習平台 | 2024.06", "受邀為行政院公務同仁錄製《洞見機遇-雲端管理與實務》大師講座數位課程，推廣公部門雲端轉型、資安韌性與現代化資訊治理思維，列入國家公務員數位學習專屬教材。", url="https://elearn.hrd.gov.tw/info/10042804", image_filename="e等公務園__洞見機遇-雲端管理與實務01.png", img_width=Inches(3.0), east_asia="Microsoft JhengHei", url_label="🔗 查證佐證出處 / URL：")
    add_media_entry(doc, "Google Cloud 官方全球客戶成功案例影音專訪 (GameSparcs APAC Customer Success Story)", "Google Cloud APAC 官方專訪 (YouTube) | 2024.03", "接受 Google Cloud 官方採訪，深入解析如何運用 GKE 多雲架構與自動擴縮容技術，支撐全球百萬級高併發手遊連線營運，創下 100% Zero Outage 零停機紀錄。", url="https://youtu.be/_kTZSZ_0lNE?si=CT2lo8c4IF0zI1Ki", east_asia="Microsoft JhengHei", url_label="🔗 官方影片連結：")

    add_subheading(doc, "2. 國際期刊論文與國家學術典藏", size=Pt(11), east_asia="Microsoft JhengHei")
    add_media_entry(doc, "Springer SCI 國際頂級期刊論文：Reversible secret-image sharing with high visual quality", "Multimedia Tools and Applications (Springer Nature, Vol. 74, Pages 10603–10626) | 2014.06", "共同作者：Ching-Chiuan Lin, Lun-Hao Liao (廖倫豪), Kuo-Feng Hwang, Shih-Chieh Chen。本論文提出基於高視覺品質陰影圖像之可逆秘密影像共享技術，在密碼學與資安影像傳輸領域具高度學術影響力。", url="https://link.springer.com/article/10.1007/s11042-012-1190-1", east_asia="Microsoft JhengHei", url_label="🔗 DOI 永久出版品連結：")
    add_media_entry(doc, "國際學術期刊論文：Utilizing GIS and GPS in Designing a Trilingual Tourist APP", "Applied Science and Management Research (Vol. 2) | 2015.05", "探討結合地理資訊系統 (GIS) 與全球定位系統 (GPS) 開發三語觀光導覽系統之行動應用架構設計與實作。", east_asia="Microsoft JhengHei")
    add_media_entry(doc, "國家圖書館博碩士學位論文法定典藏《整合系統的商業自助式入口網站植基於調適性服務導向架構資訊科技治理之研究》", "國家圖書館臺灣博碩士論文知識加值系統 (典藏代碼: 106IKTC0183002) | 2014.10", "博士論文研究：結合 ITIL V3 與 SOA 服務導向架構建立企業級自助式入口網站 IT 治理模型，並於國家圖書館永久法定典藏，擔任學術研究指導學者典藏代碼 106IKTC0183002。", url="https://ndltd.ncl.edu.tw/", east_asia="Microsoft JhengHei", url_label="🔗 國圖學術系統：")

    add_subheading(doc, "3. 主流科技媒體實名專訪與專題報導", size=Pt(11), east_asia="Microsoft JhengHei")
    add_media_entry(doc, "CIO Taiwan 官方專訪《【專訪】隆中網絡 GameSparcs IT Director Howard Liao | 善用公有雲服務 搶攻全球遊戲商機》", "CIO Taiwan 雜誌 (採訪／施鑫澤‧文／林裕洋‧刊期／2024.05) | 2024.05", "CIO Taiwan 總編輯專題專訪，實名刊登 Howard Liao, PhD 廖博士之多雲架構佈局，深入剖析採用多雲網路 (MCN) 與 Kubernetes 原生架構橫跨公有雲與邊緣運算，強化全球發行競爭力與 FinOps 降本 30%。", url="https://www.cio.com.tw/interview-howard-liao-online-gamesparcs-it-director/", image_filename="2024_CIO報導.png", img_width=Inches(3.0), east_asia="Microsoft JhengHei", url_label="🔗 媒體報導連結：")
    add_media_entry(doc, "iThome 電腦報專題企劃《【iThome 專題企劃】隆中網絡運用 MongoDB Atlas 建構遊戲資料管理平台，打造全方位休閒娛樂平台》", "iThome 電腦報官方專題企劃 | 2024.09", "技術專題企劃實名專訪廖倫豪博士，闡述以 MongoDB Atlas 建置多區域即時數據中台，兼具 99.995% 高可用性與高擴展性，成功支撐全球跨區高併發讀寫需求。", url="https://www.ithome.com.tw/pr/163534", image_filename="IThome_MangoDB.png", img_width=Inches(3.2), east_asia="Microsoft JhengHei", url_label="🔗 媒體報導連結：")
    add_media_entry(doc, "DIGITIMES 科技網報導：MongoDB 8.0 問世，助企業接軌生成式 AI (隆中網絡等企業案例)", "DIGITIMES 科技網 | 2024.09", "DIGITIMES 深度報導 MongoDB 8.0 發表會，實名收錄隆中網絡技術團隊架構分享與 AI 數據治理之成功應用案例。", url="https://www.digitimes.com.tw/tech/dt/n/shwnws.asp?id=0000704377_KI14VVDK6CQ8TB8OXC3GJ", east_asia="Microsoft JhengHei", url_label="🔗 媒體報導連結：")

    add_subheading(doc, "4. 業界獲獎榮譽與社會公益責任 (CSR)", size=Pt(11), east_asia="Microsoft JhengHei")
    add_bullet(doc, "雲端架構卓越獎 (Cloud Architecture Excellence Award) — 表彰 GKE 多雲架構、高可用性、營運韌性與 FinOps 降本 30% 之卓越成效。", bold_prefix_colon=False, east_asia="Microsoft JhengHei")
    add_bullet(doc, "受邀擔任資安高峰會 (Cybersecurity Summit) Keynote 主講人，分享混合雲資安、零信任架構、ISO 27001、ISO/IEC 42001 與營運韌性實務。", bold_prefix_colon=False, east_asia="Microsoft JhengHei")
    add_bullet(doc, "受邀於 CIO 價值學院及技術社群分享多雲治理、AI 治理、企業數位轉型與資安戰略。", bold_prefix_colon=False, east_asia="Microsoft JhengHei")
    add_media_entry(doc, "財團法人康善基金會 (Cancer NoNo Foundation) 董事 兼 資訊長 (CIO)", "財團法人康善基金會 (CSR 永續公益) | 2022 – 至今", "擔任基金會董事兼資訊長，推動兒童健康飲食教育、便當童話劇、數位公益與醫療資訊安全系統整合，落實企業社會責任 (CSR) 與數位治理。", url="https://www.cancer-nono.org.tw/web/about/page.php?lang=zh_tw&scid=7&sid=5", east_asia="Microsoft JhengHei", url_label="🔗 基金會官方連結：")

    out_file = os.path.join(target_dir, "Howard_Liao_CISO_Resume_ZH.docx")
    doc.save(out_file)
    print(f"Generated: {out_file}")

# 3. GENERATE JAPANESE (Howard_Liao_CISO_Resume_JA.docx)
def build_ja_docx():
    doc = create_base_doc()
    add_header(
        doc,
        name_text="廖倫豪 博士 (Howard Liao, Ph.D.)",
        title_text="グループ最高情報セキュリティ責任者 (Group CISO)",
        subtitle_text="グローバルサイバーセキュリティ・デジタルトラスト・レジリエンス統括 (取締役会・CEO向け 職務経歴書)",
        contact_text1="台湾 (グローバル拠点対応可能)  |  電話番号：+886-975-323161  |  Eメール：Liao.Howard@gmail.com",
        contact_text2="LinkedIn：linkedin.com/in/howardliao78  |  ポートフォリオ：https://howardliao.github.io/portfolio/",
        east_asia="Meiryo"
    )
    
    add_heading_1(doc, "エグゼクティブサマリー (Executive Profile)", east_asia="Meiryo")
    add_body_p(doc, "27年以上の企業ITリーダーシップ、15年以上のサイバーセキュリティ実務、そして上場企業・多国籍企業・規制対象業界において10年以上にわたりセキュリティ戦略、クラウドガバナンス、デジタルレジリエンス、DXを主導してきたCISOレベルのエグゼクティブ。", east_asia="Meiryo")
    add_body_p(doc, "取締役会レベルのセキュリティガバナンス力と、ゼロトラスト、IAM、マルチクラウドセキュリティ、SOC/SIEM/EDR、インシデント対応、DevSecOps、データ保護、サードパーティリスク管理、AIガバナンス(ISO 42001)における高度な技術力を兼備。", east_asia="Meiryo")
    add_body_p(doc, "基幹業務システム、機密情報、知的財産、ERP/SaaSプラットフォーム、越境データフローの保護に精通。リスクベースのアプローチにより、セキュリティ成熟度、事業継続性、監査即応性、デジタルトラストを継続的に向上。", east_asia="Meiryo")
    add_body_p(doc, "セキュリティリスクや技術課題を、取締役会が重視する事業利益、財務成果、コンプライアンス指標へと翻訳する卓越した実績。重要資産を保護しながらグローバル事業拡大を強力に後押しするスケーラブルな運用モデルを構築。", east_asia="Meiryo")
    add_body_p(doc, "IT/OT融合および多拠点製造環境での豊富な経験により、複雑なインフラ統合、サプライチェーンセキュリティ、パートナー連携においても確固たる実績を有す。", east_asia="Meiryo")

    add_heading_1(doc, "取締役会・経営陣向けリーダーシップ価値提案 (Leadership Value Proposition)", east_asia="Meiryo")
    add_bullet(doc, "CISO & 取締役会統括：リスクマトリクス、RTO/RPO、SLA、ROI/TCO、成熟度ロードマップを用いて、セキュリティ態勢とリスクを明確な事業価値へ変換。", east_asia="Meiryo")
    add_bullet(doc, "グループ統合ガバナンス：本社、海外拠点、マルチクラウド、SaaS、マネージドサービス事業者、パートナー企業を網羅するグループ標準セキュリティ運用モデルを確立。", east_asia="Meiryo")
    add_bullet(doc, "デジタルトラスト & 復原力：アイデンティティ統制、監査証跡保護、暗号化、イミュータブルバックアップ、インシデント対応体制により重要データと知財を全方位で防御。", east_asia="Meiryo")
    add_bullet(doc, "クラウドネイティブ技術深度：AWS、Azure、GCP、Kubernetes/GKE、Zero Trust、APIセキュリティ、WAF/WAAP、SIEM、EDR、IaC自動化の戦略と実践を高度に融合。", east_asia="Meiryo")
    add_bullet(doc, "AIガバナンス (ISO 42001)：ISO/IEC 42001規格に準拠し、GenAIの責任ある活用を推進しながら、Shadow AI、データ流出、IPリスク、第三者モデルリスクを徹底統制。", east_asia="Meiryo")
    add_bullet(doc, "事業成長イネーブラー：事業、財務、法務、人事、開発、運用部門と密接に連携し、セキュリティを企業成長、DX、顧客信頼の基盤として組み込む。", east_asia="Meiryo")

    add_heading_1(doc, "8大コアコンピテンシー & 技術統治体系 (Core Competencies)", east_asia="Meiryo")
    comp_data_ja = [
        ("セキュリティ戦略・ガバナンス・リスク管理 (Strategy, Governance & Risk)", [
            "グループセキュリティ戦略、運用モデル、ポリシー策定、リスク選好度、成熟度評価、KPI/KRI設計、取締役会報告、中長期ロードマップ。",
            "ISO 27001/27002、NIST CSF、Zero Trust、ベンダーリスク評価、デューデリジェンス、監査即応性、コントロールアセスメント。",
            "リスク低減、ROI/TCO、RTO/RPO、SLA、事業継続性を通じたセキュリティ投資ガバナンス。",
            "取締役会報告、外部監査対応、主要顧客セキュリティレビュー、危機管理統治。"
        ]),
        ("デジタルトラスト・データ統治・監査証跡 (Digital Trust & Auditability)", [
            "重要業務システム、機密記録、個人情報、知的財産、ワークフロー、データ基盤のセキュリティガバナンス。",
            "資産棚卸し、重要度分類、リスク評価、セキュリティベースライン、監査証跡管理、SOP策定。",
            "最小権限、職務分掌(SoD)、アクセス定期棚卸し、ログ保護、暗号化、バックアップ・復元、変更管理統制。",
            "データの分類、アクセス、保管、アーカイブ、廃棄、完全性、追跡可能性を網羅するライフサイクル管理。",
            "ERP、CRM、HRIS、財務システム、文書管理、SaaS連携におけるセキュリティアーキテクチャとベンダー保証。"
        ]),
        ("アイデンティティ・ゼロトラスト・データ保護 (Identity, Zero Trust & DLP)", [
            "企業IAM、SSO、MFA、条件付きアクセス、特権アクセス管理(PAM)、最小権限、JMLライフサイクル、サードパーティID管理。",
            "データ分類、DLP、暗号化・鍵管理、エンドポイント保護、CASB/SSE、越境データガバナンス。",
            "知的財産、機密R&D情報、顧客・パートナーデータ、財務記録、内部コア資産の完全保護。"
        ]),
        ("マルチクラウド・SaaS・エンタープライズアプリ (Multi-Cloud & AppSec)", [
            "AWS/Azure/GCPセキュリティ統治、ランディングゾーン、集中ログ、CSPM/CNAPP、ネットワーク分離、WAF/WAAP、API保護、DR。",
            "ERP/SAP、CRM、HRIS、財務アプリ、データ基盤、SaaSサービスのセキュア統制。",
            "クラウド/SaaSデューデリジェンス、責任共有モデル、アーキテクチャレビュー、継続的アシュアランス、BCP。",
            "重要サービスのマルチゾーン・マルチリージョン高可用性設計および越境ディザスタリカバリ。"
        ]),
        ("セキュリティ運用・レジリエンス・インシデント対応 (SecOps & Incident Response)", [
            "SOC/MDR運用、SIEM、EDR/XDR、集中ログ、検知エンジニアリング、脅威インテリジェンス、脆弱性管理、事故対応体制。",
            "ランサムウェア対策、イミュータブルバックアップ、ディザスタリカバリ、危機対応、机上演習、復旧統制。",
            "ELK Stack、Graylog、Prometheus、LibreNMS、クラウドネイティブ監視によるエンドツーエンドの可観測性。",
            "MTTR、検知カバー率、事故重大度、対応実効性、復旧能力などの定量的メトリクス管理。"
        ]),
        ("DevSecOps・アプリケーション・サプライチェーン (DevSecOps & Supply-Chain)", [
            "セキュアSDLC、脅威モデリング、SAST/DAST/SCA、シークレット管理、APIセキュリティ、CI/CD統制、IaC/Terraform、K8s/GKE、SBOM。",
            "ソフトウェアサプライチェーンセキュリティ、オープンソース管理、セキュアコードレビュー、脆弱性修復。",
            "社内開発チーム、外部委託先、SIer、API統合、データパイプラインの包括的セキュリティ統制。"
        ]),
        ("AIガバナンス & セキュアイノベーション (AI Governance & Innovation)", [
            "ISO/IEC 42001準拠AIガバナンス、責任あるAI方針、承認ユースケース管理、モデル・ベンダーリスク評価、ライフサイクル管理。",
            "GenAIデータ漏洩防止(DLP)、Shadow AI抑止、セキュアRAGナレッジ統治、アクセス制御、プロンプト保護、Human-in-the-Loop。",
            "AI活用によるリスクスコアリング、異常検知、脅威相関分析、セキュリティアナリティクス。"
        ]),
        ("リーダーシップ & ステークホルダーマネジメント (Leadership & Stakeholder Mgmt)", [
            "台湾、中国、アジア太平洋、欧州、グローバルマネージドパートナーにまたがる多国籍チーム統括実績。",
            "セキュリティ、IT運用、クラウド、インフラ、SRE、DevOps、アプリ開発、外部パートナーの組織マネジメント。",
            "年間数百万〜1,400万ドル規模のグローバルIT・セキュリティ予算管理。",
            "中英バイリンガルによる取締役会報告、外部監査対応、主要顧客監査対応、危機管理リーダーシップ。"
        ])
    ]
    for d_title, bullets in comp_data_ja:
        add_subheading(doc, d_title, east_asia="Meiryo")
        for b in bullets:
            add_bullet(doc, b, bold_prefix_colon=False, east_asia="Meiryo")

    add_heading_1(doc, "職務経歴 & 主要実績 (Professional Experience)", east_asia="Meiryo")
    ja_labels = {"scope": "統括範囲 (Leadership Scope)", "ach": "主要成果・定量的実績 (Selected Achievements & Impact)", "ctx": "補足情報 (Additional Context)", "leaving": "離職理由 (Reason for Leaving)"}
    
    # 1. 副社長
    add_exp_item(
        doc,
        role="副社長 兼 グループセキュリティ・DX統括 (Vice President)",
        group_note="(高雄・上海・台北 / 盛欣・盛碁ネットワーク / 中国Poker City / 台北Finger International) 上場関連グループ",
        company_line="多国籍インターネット・デジタルプラットフォームグループ | 台湾・海外拠点",
        period="2025年5月 – 現在",
        desc="CISOレベルの役割として、複数事業部門および海外拠点におけるグループ全体のサイバーセキュリティ戦略、マルチクラウド統治、レジリエンス、AIガバナンスを統括。",
        scopes=[
            "複数事業部門および国内外拠点にまたがるグループ全体のIT・セキュリティガバナンス。",
            "セキュリティ、インフラ、アプリケーション、運用、外部パートナーからなる約35〜40名の専門組織を統括。",
            "年間約1,000万〜1,200万米ドルのIT・セキュリティ予算の統括管理。"
        ],
        achs=[
            "ISO 27001およびNIST CSFに基づくグループ共通のセキュリティ統治体制を確立し、取締役会レベルの監視へ格上げ。",
            "資産重要度、脅威情報、脆弱性、統制成熟度、サードパーティリスクを網羅する企業リスク管理システムと定量マトリクスを構築。",
            "年度予算、OGSM目標、事業優先度、定量的リスク低減効果と連動した複数年セキュリティロードマップを策定。",
            "基幹アプリ、データ基盤、パートナー連携を対象に、ゼロトラスト原則に準拠したID・エンドポイント・クラウドアクセスアーキテクチャを設計。",
            "集中SOC/SIEM、EDR、ログ統合、監視・対応SOPを導入し、重大セキュリティインシデントを約30%削減、検知・初動時間を大幅短縮。",
            "マルチクラウド可観測性および設定遵守ダッシュボードを構築し、経営陣の可視性、監査対応、顧客DD対応を効率化。",
            "ISO/IEC 42001に基づくAIガバナンスを導入し、利用規程、リスク評価、データ保護、モデル管理、外部AIサービス統制を確立。",
            "ROI/TCO、RTO/RPO、SLA、リスク低減効果の精緻な分析により、ゼロトラスト、SIEM、DR、クラウドセキュリティへの役員投資承認を獲得。",
            "IT、財務、法務、人事、開発、事業部門、外部パートナー間の緊密な連携を推進し、全社DXプロジェクトを完遂。"
        ],
        leaving="長期的なキャリア形成、家族との生活設計、およびグループ全体を見据えた最高セキュリティ責任者(CISO)職への就任を見据え、台湾北部拠点のポジションを希望。",
        east_asia="Meiryo",
        labels=ja_labels
    )

    # 2. GameSparcs
    add_exp_item(
        doc,
        role="IT Director (情報技術統括部長)",
        group_note="Longzhong Network Co., Ltd. / GameSparcs (台中本社、台北・ロサンゼルス・シドニー・マルタ・杭州・成都拠点 / 台湾上場企業)",
        company_line="グローバルオンラインゲーム・エンターテインメントプラットフォーム上場企業 | 台中, 台湾",
        period="2022年9月 – 2025年4月",
        desc="世界数百万のユーザーを支える上場プラットフォームにおいて、グローバルクラウドアーキテクチャ、セキュリティ統治、プラットフォーム復原力、DevOps生産性、高可用性運用を主導。",
        scopes=[
            "クラウド運用、DevOps、SRE、セキュリティ、インフラ、アプリサポートからなる約30〜35名のエンジニア組織を統括。",
            "年間約1,200万〜1,400万米ドルのマルチクラウドインフラおよびセキュリティ予算を管理。",
            "24時間365日無停止、大規模トランザクション処理、高耐久性が求められる高負荷プラットフォームを支援。"
        ],
        achs=[
            "GKEを中心とするマルチゾーン・マルチリージョン構成と越境DRを構築し、24/7稼働と基幹サービス障害ゼロ(100% Zero Outage)を達成。",
            "集中可観測性、APM、SIEM、ログ分析基盤を導入し、リアルタイムイベント相関分析によりMTTRを約30%短縮。",
            "ISO 27001およびNIST CSFに準拠したIAM、資産管理、セキュア開発、クラウド運用、インシデント対応ポリシーを全面適用。",
            "IAM、SSO、MFA、最小権限、セキュアリモートアクセスによるゼロトラスト原則を導入し、分散拠点と特権ユーザーの防御を強化。",
            "APIセキュリティ、WAF/WAAP、DDoS防御、アンチボット、レート制限、認可制御により外部アタックサーフェスを強固に防御。",
            "インシデント対応プレイブック、机上演習、Red/Blue演習、危機管理手順を確立し、組織の対応力強化と被害極小化を実現。",
            "SAST/DAST、シークレット管理、CI/CDセキュリティ、IaC統治、コンテナセキュリティ、SBOMを開発・デプロイ工程へ統合。",
            "マルチクラウドFinOpsを推進し、トラフィック急増時にも高パフォーマンスを維持しつつ、クラウドコストの約30%最適化を達成。",
            "可用性・復原力・コスト最適化を高度に融合したGKEマルチクラウド基盤が高く評価され、「クラウドアーキテクチャ卓越賞」を受賞。"
        ],
        leaving="グループ全体へのより広範なガバナンス責任を担うシニアセキュリティリーダーシップポジションへの挑戦のため。",
        east_asia="Meiryo",
        labels=ja_labels
    )

    # 3. Hongen
    add_exp_item(
        doc,
        role="IT Director (情報システム責任者)",
        group_note="Hongen Technology (深セン本社、新竹・杭州・成都拠点 / 泓晏科技)",
        company_line="電子機器製造 & ハイテクサービス企業 | 台湾・中国",
        period="2018年3月 – 2022年8月",
        desc="台湾および中国本土の複数拠点にまたがる企業IT、セキュリティ、データガバナンス、DX、工場インフラ運用を統括。",
        scopes=[
            "IT、セキュリティ、アプリ、インフラ、製造ITからなる約30〜40名の専門組織を統括。",
            "年間約500万〜700万米ドルのIT・セキュリティ予算を執行。",
            "多拠点製造、R&D、サプライチェーン、生産管理、財務、調達環境を全面支援。"
        ],
        achs=[
            "オフィスおよび製造現場を対象にISO 27001準拠のセキュリティガバナンス、リスク評価、資産管理、IAM、ネットワーク分離、監査性を確立。",
            "ERP、PLM、SCM、ワークフロー、分析基盤のセキュアアーキテクチャを設計し、可用性、追跡可能性、事業継続性を強化。",
            "ERP、MES、PLM、APS、HR基盤のセキュア統合を主導し、認証・暗号化・監査ログを伴う安全なデータ連携を実現。",
            "データガバナンスとBI基盤を構築し、在庫最適化、異常検知、経営意思決定、M&A後のデータ統合を支援。",
            "セキュリティ強化と工場連続稼働を両立するリスクベースのパッチ管理、インシデント対応、BCP手順を策定。",
            "台湾・中国拠点のIT、製造、調達、財務部門および外部パートナー間のセキュリティプロジェクトを統括。",
            "中国大手製造グループによる買収時において、IT組織統合とデータ資産デューデリジェンスを主導。"
        ],
        add_ctx="豊富なIT/OT融合経験により、複雑な産業インフラ、サプライチェーン、パートナー統合にも即応可能。",
        leaving="大手製造グループへの統合完了後、更なる広範なセキュリティおよびテクノロジーリーダーシップ機会を追求するため。",
        east_asia="Meiryo",
        labels=ja_labels
    )

    # 4. HyWeb
    add_exp_item(
        doc,
        role="IT Manager (IT部門マネージャー)",
        group_note="HyWeb Technology Co., Ltd. (新竹本社、台北・台中・高雄・タイ・北京拠点 / 台湾上場企業)",
        company_line="上場大手ITサービス企業 | 台湾・中国・タイ",
        period="2014年5月 – 2018年2月",
        desc="台北、新竹、台中、北京、タイ拠点における地域IT運用、ITSMサービス管理、データセンター戦略、クロスボーダー技術支援を統括。",
        scopes=[
            "インフラ、基幹アプリ、サービスデスク、運用保守からなる約34〜45名の技術組織をマネジメント。",
            "地域IT予算、データセンター運用、バックアップ構成、ディザスタリカバリ計画を主導。"
        ],
        achs=[
            "ITSMおよび地域共通サービス管理を導入し、アジア太平洋各拠点のサービス品質、透明性、インシデント対応力を向上。",
            "地域データセンター戦略、高可用性構成、バックアップ統治、DR計画を策定し基幹サービスを安定稼働。",
            "PMPおよびアジャイル手法を導入し、国境を越えたプロジェクト推進力と納期予見性を強化。",
            "分散拠点におけるインフラ標準化、セキュリティ運用規範、ITサービス体制を確立。"
        ],
        east_asia="Meiryo",
        labels=ja_labels
    )

    # 5. Kuang Nan
    add_exp_item(
        doc,
        role="IT Manager / DX推進責任者",
        group_note="Kuang Nan Group / Fengchen Group – Aiju Computer (台中本社、台北・深セン・上海・張家港拠点 / 上場企業)",
        company_line="小売・流通・サプライチェーン & ITサービスグループ | 台湾・中国",
        period="2011年7月 – 2014年4月",
        desc="M&A局面における基幹システム統合、ビジネスデータガバナンス、全社デジタル変革を主導。",
        scopes=[
            "15〜20名のITエンジニア組織をマネジメント。",
            "ERP、POS、CRM、データ統合、事業継続性、アプリケーション統治予算を担当。"
        ],
        achs=[
            "ERP、POS、CRM、分析基盤を統合し、全店舗の売上、在庫回転、顧客動態のリアルタイム可視化を実現。",
            "越境M&Aに伴うデータクレンジング、マスターデータ管理(MDM)、システム移行を指揮し業務継続を担保。",
            "統合データ基盤により、経営陣の意思決定を支援するリアルタイムダッシュボードを構築。"
        ],
        east_asia="Meiryo",
        labels=ja_labels
    )

    # 6. Borland
    add_exp_item(
        doc,
        role="IT Manager",
        group_note="Borland Taiwan (米商ポーランド)",
        company_line="世界的ソフトウェア & アプリケーション開発ベンダー | 台湾・中国",
        period="2008年3月 – 2011年6月",
        desc="台湾・中国のR&Dおよび事業組織向けに、越境IT協調基盤、高可用性基幹システム、データ保護環境を構築。",
        scopes=None,
        achs=[
            "ミッションクリティカルなトランザクション環境向けに高可用性DB、アプリ基盤、バックアップ・DR構成を設計。",
            "ソフトウェアライフサイクル管理(ALM)ツールおよび開発コラボレーション基盤の全社展開を完遂。",
            "多国籍開発環境において最高水準のシステム安定性と越境データ保護を維持。"
        ],
        east_asia="Meiryo",
        labels=ja_labels
    )

    # 7. Sybase
    add_exp_item(
        doc,
        role="IT Manager",
        group_note="Sybase (米商サイベース)",
        company_line="世界的データベース & エンタープライズソフトウェアベンダー | アジア太平洋",
        period="2002年8月 – 2008年2月",
        desc="アジア太平洋地域のデータセンター統合、インフラ標準化、高可用性DB運用、通信・金融顧客向けミッションクリティカル支援を統括。",
        scopes=[
            "アジア太平洋地域のインフラチーム約10〜15名のエンジニアを統括。",
            "大規模データベース、クラスタリング、フェイルオーバー、データセンター運用を支援。"
        ],
        achs=[
            "APAC全域で統一Active Directoryの構築、データセンター統合、高可用性DBフェイルオーバー戦略を実行。",
            "通信事業者および金融機関向けにキャリアグレードのDBチューニング、高可用性、リカバリソリューションを提供。",
            "グローバル組織統合および顧客プロジェクトにおいて、100%の業務継続と監査コンプライアンスを達成。"
        ],
        east_asia="Meiryo",
        labels=ja_labels
    )

    add_heading_1(doc, "保有プロフェッショナル資格 (Certifications)", east_asia="Meiryo")
    for cert in [
        "ISO 27001 情報セキュリティマネジメントシステム — リード審査員 / 内部監査員 (Lead Auditor)",
        "ISO/IEC 42001 人工知能マネジメントシステム (AIMS) — AIガバナンスリード審査員",
        "Project Management Professional — PMP 認定国際プロジェクトマネージャー",
        "Certified ScrumMaster — CSM 認定スクラムマスター",
        "Oracle Certified Professional — OCP 認定プロフェッショナル",
        "Oracle Certified Associate — OCA 認定アソシエイト",
        "ESG サステナビリティプランナー (ESG Sustainability Planner)"
    ]:
        add_bullet(doc, cert, bold_prefix_colon=False, east_asia="Meiryo")

    add_heading_1(doc, "学歴・学位 (Academic Education)", east_asia="Meiryo")
    edu_list_ja = [
        ("情報技術管理 博士 (Ph.D. in IT Management)", "朝陽科技大学 情報技術管理研究所 (College of Informatics)\n2009年9月 – 2013年6月", "研究分野：サイバーセキュリティ統治、クラウドセキュリティ、AIガバナンス、データガバナンス、ソフトウェア工学、DevSecOps、DX戦略。"),
        ("情報科学 修士 (Master of Science in Information Science)", "国立中興大学 情報科学工学研究所\n2004年9月 – 2006年6月", "研究分野：現代暗号学 (Modern Cryptography)、データマイニング、データ統治、ソフトウェア工学、システム分析、企業システム設計。"),
        ("情報科学工学 学士 (Bachelor of Computer Science)", "国立中興大学 情報科学工学科\n2000年9月 – 2004年6月", "専攻分野：ソフトウェア開発ライフサイクル(SDLC)、ソフトウェア工学、システム分析、エンタープライズアーキテクチャ。")
    ]
    for d, s, desc in edu_list_ja:
        add_subheading(doc, d, size=Pt(10.5), east_asia="Meiryo")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(s)
        set_run_font(r, size=Pt(9.5), bold=True, color=COLOR_SECONDARY, east_asia="Meiryo")
        add_body_p(doc, desc, space_after=4, east_asia="Meiryo")

    add_heading_1(doc, "受賞歴・基調講演・学術論文・メディア掲載・CSR", east_asia="Meiryo")
    
    add_subheading(doc, "1. 基調講演・講義実績 (Keynotes & Speaking)", size=Pt(11), east_asia="Meiryo")
    add_media_entry(doc, "MongoDB.local Taipei 基調講演スピーカー", "MongoDB Inc. 主催カンファレンス | 2024年9月", "Keynoteスピーカーとして登壇。MongoDB AtlasフルマネージドDBの高拡張性と99.995%可用性を活かし、世界百万人規模のリアルタイム対戦を支える高並行アーキテクチャを発表。", url="https://www.ithome.com.tw/pr/163534", image_filename="2024_MongoDB02.jpg", img_width=Inches(3.4), east_asia="Meiryo", url_label="🔗 検証済み公式URL：")
    add_media_entry(doc, "CIO Taiwan Value Academy 第17回大会 基調講演", "CIO Taiwan Value Academy | 2024年5月", "上場企業CIO/CISOを対象に、マルチクラウド資源統治、セキュリティコンプライアンス、FinOpsコスト最適化、レジリエンス戦略を講演。", url="https://www.cio.com.tw/events/value-index-2/", image_filename="問題比答案重要.png", img_width=Inches(3.2), east_asia="Meiryo", url_label="🔗 検証済み公式URL：")
    add_media_entry(doc, "行政院人事行政総処 e等公務園 マスター講座", "行政院人事行政総処 デジタル学習基盤 | 2024年6月", "公務員向けマスター講座講師として、公共部門のクラウドトランスフォーメーション、ゼロトラスト、現代的ITガバナンス講義を担当。", url="https://elearn.hrd.gov.tw/info/10042804", image_filename="e等公務園__洞見機遇-雲端管理與實務01.png", img_width=Inches(3.0), east_asia="Meiryo", url_label="🔗 検証済み公式URL：")
    add_media_entry(doc, "Google Cloud 公式グローバル導入事例 (ビデオ取材)", "Google Cloud APAC 公式ビデオ取材 | 2024年3月", "Google Cloud公式取材において、GKE自動スケーリングと高可用構成により、ピーク時にも障害ゼロ(100% Zero Outage)を達成したアーキテクチャを解説。", url="https://youtu.be/_kTZSZ_0lNE?si=CT2lo8c4IF0zI1Ki", east_asia="Meiryo", url_label="🔗 公式動画URL：")

    add_subheading(doc, "2. 国際学術論文 & 国家学術アーカイブ", size=Pt(11), east_asia="Meiryo")
    add_media_entry(doc, "Springer SCI 国際トップ学術ジャーナル論文掲載: Reversible secret-image sharing with high visual quality", "Multimedia Tools and Applications (Springer Nature, Vol. 74) | 2014年6月", "著者：Ching-Chiuan Lin, Lun-Hao Liao (廖倫豪), Kuo-Feng Hwang, Shih-Chieh Chen。高品質なシャドウ画像を用いた可逆秘密画像共有技術を提案し、暗号セキュリティ分野で高インパクトを記録。", url="https://link.springer.com/article/10.1007/s11042-012-1190-1", east_asia="Meiryo", url_label="🔗 DOI 永久論文リンク：")
    add_media_entry(doc, "国際学術論文発表 (Trilingual Tourist APP)", "Applied Science and Management Research (Vol. 2) | 2015年5月", "GISとGPSを融合させた多言語観光プラットフォームのモバイルシステムアーキテクチャ設計と実装に関する研究。", east_asia="Meiryo")
    add_media_entry(doc, "国家図書館 博士学位論文 法定所蔵 & 指導教授アーカイブ", "国家図書館 台湾博修士論文アーカイブ (所蔵コード: 106IKTC0183002) | 2014年10月", "ITIL V3とSOAを融合したエンタープライズポータルITガバナンスモデルを確立。国家図書館にて永久所蔵、学術指導学者コード 106IKTC0183002 として登録。", url="https://ndltd.ncl.edu.tw/", east_asia="Meiryo", url_label="🔗 国家図書館公式URL：")

    add_subheading(doc, "3. メディア取材・報道実績", size=Pt(11), east_asia="Meiryo")
    add_media_entry(doc, "CIO Taiwan 誌 カバー特集 実名独占インタビュー", "CIO Taiwan 誌 (編集長インタビュー / 2024.05刊期) | 2024年5月", "編集長による特別インタビュー。マルチクラウドネットワーキング(MCN)、Kubernetesネイティブ設計、FinOpsによる30%コスト削減戦略を実名公開。", url="https://www.cio.com.tw/interview-howard-liao-online-gamesparcs-it-director/", image_filename="2024_CIO報導.png", img_width=Inches(3.0), east_asia="Meiryo", url_label="🔗 取材記事URL：")
    add_media_entry(doc, "iThome コンピュータ報 専門企画 実名特集", "iThome コンピュータ報 専門企画 | 2024年9月", "技術特集取材において、99.995%の可用性を誇るMongoDB Atlasマルチリージョンリアルタイムデータハブの設計思想を解説。", url="https://www.ithome.com.tw/pr/163534", image_filename="IThome_MangoDB.png", img_width=Inches(3.2), east_asia="Meiryo", url_label="🔗 特集記事URL：")
    add_media_entry(doc, "DIGITIMES 科技網 報道: MongoDB 8.0 発表", "DIGITIMES 科技網 | 2024年9月", "MongoDB 8.0発表会報道において、GameSparcsのAIデータ基盤および高並行処理アーキテクチャ事例が紹介。", url="https://www.digitimes.com.tw/tech/dt/n/shwnws.asp?id=0000704377_KI14VVDK6CQ8TB8OXC3GJ", east_asia="Meiryo", url_label="🔗 報道URL：")

    add_subheading(doc, "4. 受賞歴 & 企業の社会的責任 (CSR)", size=Pt(11), east_asia="Meiryo")
    add_bullet(doc, "クラウドアーキテクチャ卓越賞 (Cloud Architecture Excellence Award) — 百万人規模のトラフィックを支え、100% Zero Outageと約30%のクラウドコスト削減を両立した功績により受賞。", bold_prefix_colon=False, east_asia="Meiryo")
    add_bullet(doc, "サイバーセキュリティサミットにて基調講演スピーカーを務め、ハイブリッドクラウド、ゼロトラスト、ISO 27001、ISO 42001の実践を共有。", bold_prefix_colon=False, east_asia="Meiryo")
    add_bullet(doc, "CIOアカデミーおよび技術コミュニティにてマルチクラウド統治、AIガバナンス、DX戦略を講演。", bold_prefix_colon=False, east_asia="Meiryo")
    add_media_entry(doc, "財団法人康善基金会 (Cancer NoNo Foundation) 理事 兼 CIO", "財団法人康善基金会 (CSR活動) | 2022年 – 現在", "理事兼CIOとして、子どもの健康的な食育啓発、医療情報システムのセキュア連携、デジタル公益活動を推進し、企業の社会的責任(CSR)を実践。", url="https://www.cancer-nono.org.tw/web/about/page.php?lang=zh_tw&scid=7&sid=5", east_asia="Meiryo", url_label="🔗 基金会公式URL：")

    out_file = os.path.join(target_dir, "Howard_Liao_CISO_Resume_JA.docx")
    doc.save(out_file)
    print(f"Generated: {out_file}")

if __name__ == "__main__":
    build_en_docx()
    build_zh_docx()
    build_ja_docx()
    print("All 3 trilingual resumes successfully built!")
